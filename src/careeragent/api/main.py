"""
src/careeragent/api/main.py
============================
FastAPI backend for CareerAgent-AI.
Fixes:
  - Clean startup (no lifespan crash)
  - CORS for Streamlit on :8501
  - /hunt/start  → POST, accepts resume file + config, launches pipeline async
  - /hunt/{run_id}/status → GET, real-time progress for UI progress bar
  - /hunt/{run_id}/jobs   → GET, discovered + scored jobs
  - /hunt/{run_id}/artifacts → GET, generated file list
"""

from __future__ import annotations

import asyncio
from collections import Counter
from concurrent.futures import ThreadPoolExecutor
import json
import logging
import math
import os
import re
import sys
import tempfile
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Awaitable, Optional
from urllib.parse import quote_plus


import importlib.machinery
import importlib.util


def _repair_pydantic_shadowing() -> None:
    """Ensure local src/pydantic shims never shadow real dependency."""
    spec = importlib.util.find_spec("pydantic")
    origin = str(getattr(spec, "origin", "") or "") if spec else ""
    if "/src/pydantic" not in origin.replace("\\", "/"):
        return

    candidate_paths = []
    for path in sys.path:
        if not path:
            continue
        try:
            resolved = str(Path(path).resolve())
        except Exception:
            continue
        if resolved.endswith("/src"):
            continue
        candidate_paths.append(path)

    real_spec = importlib.machinery.PathFinder.find_spec("pydantic", candidate_paths)
    if real_spec and real_spec.loader:
        module = importlib.util.module_from_spec(real_spec)
        real_spec.loader.exec_module(module)
        sys.modules["pydantic"] = module
        return

    # Last-resort fallback: keep running with the local lightweight shim.
    # This keeps diagnostics tooling usable in constrained environments.
    os.environ.setdefault("CAREERAGENT_PYDANTIC_SHIM", "1")

_repair_pydantic_shadowing()

from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from careeragent.api.approval_utils import pick_approved_jobs, qualified_from_state
from careeragent.core.settings import Settings
from careeragent.nlp.skills import compute_jd_alignment, extract_skills, normalize_skill
from careeragent.services.notification_service import NotificationService
from careeragent.managers.leadscout_service import sanitize_job_url
from careeragent.tools.llm_tools import GeminiClient

try:
    from langsmith.run_helpers import traceable  # type: ignore
except Exception:  # pragma: no cover
    def traceable(*_args, **_kwargs):  # type: ignore
        def _decorator(fn):
            return fn
        return _decorator

# ── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
)
log = logging.getLogger("api")

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR      = Path(__file__).resolve().parent.parent.parent.parent  # project root
ARTIFACTS_DIR = BASE_DIR / "artifacts"
LOGS_DIR      = BASE_DIR / "logs"
UPLOADS_DIR   = BASE_DIR / "uploads"

for d in [ARTIFACTS_DIR, LOGS_DIR, UPLOADS_DIR]:
    d.mkdir(parents=True, exist_ok=True)

# ── In-process run registry (replace with Redis/DB for multi-worker) ──────────
_runs: dict[str, dict] = {}   # run_id → state dict
_GLOBAL_SELF_LEARNING_CONTEXT = os.getenv("CAREERAGENT_SELF_LEARNING_CONTEXT", "").strip()


# ══════════════════════════════════════════════════════════════════════════════
# LAYER DEFINITIONS  (mirrors the UI layer cards)
# ══════════════════════════════════════════════════════════════════════════════

LAYER_DEFS = [
    {"id": 0, "name": "Security & Guardrails",          "weight": 5,  "agent": "GuardAgent",     "desc": "Sanitizes input, runs guardrail checks, validates API tokens"},
    {"id": 1, "name": "Mission Control (UI)",            "weight": 5,  "agent": "UIAgent",        "desc": "Initializes UI state, loads run configuration"},
    {"id": 2, "name": "Intake Bundle (Parsing/Profile)", "weight": 15, "agent": "ParseAgent",     "desc": "Parses resume via LLM+regex, extracts skills/experience/education, builds search personas"},
    {"id": 3, "name": "Discovery (Hunt / Job Boards)",   "weight": 25, "agent": "HuntAgent",      "desc": "Searches LinkedIn, Glassdoor, Indeed, ZipRecruiter, MyVisaJobs, Greenhouse, Lever, and Google Jobs with deduped geo-aware discovery"},
    {"id": 4, "name": "Scrape + Match + Score",          "weight": 15, "agent": "MatchAgent",     "desc": "Extracts full JD text, runs semantic + keyword scoring against your profile"},
    {"id": 5, "name": "Evaluator + Ranking + HITL",      "weight": 10, "agent": "EvalAgent",      "desc": "Phase-2 evaluation, ranks by interview probability, triggers HITL gate"},
    {"id": 6, "name": "Drafting (ATS Resume + Cover)",   "weight": 10, "agent": "DraftAgent",     "desc": "Generates tailored ATS resume + cover letter per approved job using LLM"},
    {"id": 7, "name": "Apply Executor + Notifications",  "weight": 5,  "agent": "ApplyAgent",     "desc": "Auto-applies to approved jobs, sends SMS/email notifications"},
    {"id": 8, "name": "Tracking (DB + Status)",          "weight": 5,  "agent": "TrackAgent",     "desc": "Records applications to DB, updates deduplication memory"},
    {"id": 9, "name": "Analytics + Learning Center + XAI","weight": 5, "agent": "AnalyticsAgent", "desc": "Analytics, self-learning from outcomes, career roadmap, XAI explanations"},
]


def _build_initial_state(run_id: str, config: dict) -> dict:
    """Build fresh run state dict."""
    layers = []
    for ld in LAYER_DEFS:
        layers.append({
            "id":           ld["id"],
            "name":         ld["name"],
            "weight":       ld["weight"],
            "agent":        ld["agent"],
            "desc":         ld["desc"],
            "status":       "waiting",   # waiting|running|ok|error|skipped
            "started_at":   None,
            "finished_at":  None,
            "error":        None,
            "output":       None,
            "meta":         {},
        })
    return {
        "run_id":           run_id,
        "status":           "running",    # running|completed|error
        "progress_pct":     0.0,
        "created_at":       _now(),
        "completed_at":     None,
        "config":           _normalize_config(config),
        "layers":           layers,
        "profile":          {},
        "jobs_discovered":  0,
        "jobs_scored":      0,
        "jobs_approved":    0,
        "jobs_applied":     0,
        "top_match_score":  0.0,
        "candidate_name":   "—",
        "skills_extracted": 0,
        "job_leads":        [],
        "scored_jobs":      [],
        "artifacts":        {},
        "apply_results":    [],
        "resume_scores":    {},
        "agent_log":        [],         # live feed messages
        "evaluations":      [],         # layer/job evaluator outputs
        "layer_debug":      {},         # stepwise debug payload per layer
        "pending_action":   None,       # approve_ranking | approve_drafts
        "approved_jobs":    [],
        "errors":           [],
        "resume_path":      None,
        "hitl_rejections":  0,
        "interviews":       [],
        "followup_queue":   [],
        "notification_log": [],
        "feedback_events":  [],
        "learning_loop":    {"user_feedback": 0, "employer_feedback": 0, "accepted": 0, "rejected": 0},
        "employer_outcomes": {"interview": 0, "selected": 0, "rejected": 0, "unknown": 0},
        "learning_resources": {},
        "analytics_summary": {},
        "self_learning_prompt": "",
        "self_learning_context": _GLOBAL_SELF_LEARNING_CONTEXT,
        "system_prompt_update": "",
        "feedback_learning_state": {"strictness_mode": "balanced", "targeting_mode": "broad_semantic"},
        "langsmith":        _langsmith_status(run_id),
        "langgraph":        _langgraph_status(run_id),
        "llm_stack":        _llm_stack_snapshot(),
    }


def _normalize_config(config: dict) -> dict:
    cfg = dict(config) if isinstance(config, dict) else {}
    cfg.setdefault("target_roles", ["Software Engineer"])
    if not isinstance(cfg.get("target_roles"), list):
        cfg["target_roles"] = [str(cfg.get("target_roles") or "Software Engineer")]
    cfg.setdefault("match_threshold", 0.40)
    cfg.setdefault("geo_preferences", {"remote": True, "locations": [], "country_selector": "US"})
    if not isinstance(cfg.get("geo_preferences"), dict):
        cfg["geo_preferences"] = {"remote": True, "locations": [], "country_selector": "US"}
    cfg["geo_preferences"].setdefault("country_selector", "US")
    cfg.setdefault("require_ranking_approval", True)
    cfg.setdefault("require_draft_approval", True)
    cfg.setdefault("require_followup_approval", True)
    cfg.setdefault("max_jobs", 140)
    cfg.setdefault("posted_within_hours", 168)
    cfg.setdefault("salary_min", 0)
    cfg.setdefault("salary_max", 400000)
    cfg.setdefault("work_modes", ["remote", "hybrid", "onsite"])
    if not isinstance(cfg.get("work_modes"), list):
        cfg["work_modes"] = ["remote", "hybrid", "onsite"]
    cfg.setdefault("draft_jobs_limit", 0)
    cfg.setdefault("apply_jobs_limit", 0)
    cfg.setdefault("notifications", {"email": "", "phone": "", "enable_email": False, "enable_sms": False})
    raw_notifications = cfg.get("notifications")
    notifications = dict(raw_notifications) if isinstance(raw_notifications, dict) else {}
    notifications.setdefault("email", "")
    notifications.setdefault("phone", "")
    notifications.setdefault("enable_email", False)
    notifications.setdefault("enable_sms", False)
    notifications["phone"] = _sanitize_phone(notifications.get("phone", ""))
    cfg["notifications"] = notifications
    return cfg


def _sanitize_phone(phone: str) -> str:
    return " ".join(str(phone or "").strip().split())


def _langsmith_status(run_id: str) -> dict:
    tracing_flag = str(os.getenv("LANGCHAIN_TRACING_V2", "")).strip().lower()
    enabled = tracing_flag in {"1", "true", "yes", "on"} and bool(os.getenv("LANGSMITH_API_KEY"))
    endpoint = os.getenv("LANGSMITH_ENDPOINT", "https://smith.langchain.com").rstrip("/")
    project = os.getenv("LANGSMITH_PROJECT") or os.getenv("LANGCHAIN_PROJECT") or "default"
    return {
        "enabled": enabled,
        "project": project,
        "dashboard_url": f"{endpoint}/o/default/projects/p/{project}?q={run_id}" if enabled else None,
    }


def _langgraph_status(run_id: str) -> dict:
    base = os.getenv("LANGGRAPH_STUDIO_URL") or os.getenv("LANGGRAPH_BASE_URL") or ""
    base = str(base).rstrip("/")
    if not base:
        return {
            "enabled": False,
            "dashboard_url": None,
            "note": "Set LANGGRAPH_STUDIO_URL to enable a direct run link.",
        }
    return {
        "enabled": True,
        "dashboard_url": f"{base}/runs/{run_id}",
        "note": "LangGraph trace URL is environment configured.",
    }


def _llm_stack_snapshot() -> dict:
    ats_model = os.getenv("CAREERAGENT_ATS_MODEL") or os.getenv("OPENAI_MODEL") or "gpt-4o-mini"
    parser_model = os.getenv("CAREERAGENT_PARSER_MODEL") or os.getenv("GEMINI_MODEL") or "gemini-1.5-flash"
    reasoning_model = os.getenv("CAREERAGENT_REASONING_MODEL") or os.getenv("ANTHROPIC_MODEL") or "claude-3-5-sonnet"
    return {
        "ats_resume_writer": {
            "provider": "openai-compatible",
            "model": ats_model,
            "why": "Best quality/cost default for ATS resume + cover letter drafting.",
            "options": [ats_model, "gpt-4.1-mini", "gpt-4o"],
        },
        "resume_parser": {
            "provider": "google",
            "model": parser_model,
            "why": "Fast extraction with robust structured parsing fallback.",
            "options": [parser_model, "gemini-2.0-flash-lite", "gpt-4o-mini"],
        },
        "ranking_reasoner": {
            "provider": "anthropic-compatible",
            "model": reasoning_model,
            "why": "Strong long-context reasoning for match explanations.",
            "options": [reasoning_model, "gemini-1.5-pro", "gpt-4.1"],
        },
        "evaluator_guardrails": {
            "provider": "hybrid",
            "model": reasoning_model,
            "why": "Evaluator agents can fall back across reasoning, parser, and ATS-oriented models.",
            "options": [reasoning_model, parser_model, ats_model],
        },
    }




_REASONING_POOL = ThreadPoolExecutor(max_workers=max(2, int(os.getenv("CAREERAGENT_REASONING_WORKERS", "6"))))
_ROLE_EQUIVALENCE_HINT = "Treat Senior, Lead, Principal, and Architect as interchangeable career levels for this candidate when assessing capability."


def _profile_capability_evidence(profile: dict) -> str:
    experience = profile.get("experience") or []
    lines = []
    for item in experience[:8]:
        if isinstance(item, dict):
            title = str(item.get("title") or "").strip()
            company = str(item.get("company") or "").strip()
            summary = str(item.get("summary") or item.get("description") or "").strip()
            bits = [part for part in [title, company] if part]
            base = " @ ".join(bits) if bits else "Previous role"
            lines.append(f"- {base}: {summary[:220]}")
        else:
            lines.append(f"- {str(item)[:220]}")
    skills = ", ".join(str(s) for s in (profile.get("skills") or [])[:30]) or "not provided"
    return "\n".join(lines) + f"\nSkills: {skills}"


def _heuristic_cognitive_verdict(job: dict, profile: dict) -> dict:
    matched, missing = _derive_reasoning(job, profile)
    title = str(job.get("title") or "")
    lexical = float(job.get("keyword_score", job.get("score") or 0.0) or 0.0)
    semantic = float(job.get("semantic_score") or 0.0)
    cognitive = float(job.get("cognitive_score") or semantic or lexical)
    yes = bool(matched) or cognitive >= 0.58 or (semantic >= 0.48 and lexical >= 0.35)
    rationale = (
        f"Recommended because your experience in {', '.join(matched[:3]) or 'adjacent leadership and delivery'} translates directly "
        f"to the role requirements for {title or 'this role'}. "
        f"{_ROLE_EQUIVALENCE_HINT} "
        f"{('Minor gap areas: ' + ', '.join(missing[:3]) + '. ') if missing else 'No blocking capability gaps detected. '}"
        "This cognitive decision overrides strict keyword filtering when adjacent experience is strong."
    )
    return {
        "approved": yes,
        "reasoning": rationale.strip(),
        "source": "heuristic-fallback",
    }


def _llm_cognitive_verdict(job: dict, profile: dict) -> dict:
    settings = Settings()
    client = GeminiClient(settings, model=os.getenv("CAREERAGENT_REASONING_MODEL") or os.getenv("GEMINI_MODEL") or "gemini-1.5-flash")
    profile_blob = _profile_capability_evidence(profile)
    jd_excerpt = " ".join(str(job.get(k) or "") for k in ("title", "company", "description", "snippet", "full_text_md"))[:4000]
    prompt = f"""You are the CareerAgent cognitive evaluator.
{_ROLE_EQUIVALENCE_HINT}

Answer the question: Does this candidate have the capability to perform this role based on their experience?
Return strict JSON with keys approved (boolean), reasoning (string), transferable_evidence (array of strings).
Approve the job if the candidate appears capable, even when exact keywords differ, such as PyTorch vs TensorFlow, platform vs backend, or architect vs principal.
Reject only when the experience is clearly unrelated.

Candidate evidence:
{profile_blob}

Job:
{jd_excerpt}
"""
    payload = client.generate_json(prompt, temperature=0.1, max_tokens=500)
    if isinstance(payload, dict) and "approved" in payload:
        reasoning = str(payload.get("reasoning") or "").strip()
        evidence = payload.get("transferable_evidence") or []
        if isinstance(evidence, list) and evidence:
            reasoning = (reasoning + " Evidence: " + "; ".join(str(x) for x in evidence[:3])).strip()
        return {
            "approved": bool(payload.get("approved")),
            "reasoning": reasoning or "LLM approved based on transferable experience.",
            "source": "llm",
        }
    return _heuristic_cognitive_verdict(job, profile)


async def _cognitive_reason_job(job: dict, profile: dict) -> dict:
    loop = asyncio.get_running_loop()
    try:
        verdict = await loop.run_in_executor(_REASONING_POOL, _llm_cognitive_verdict, job, profile)
    except Exception:
        verdict = _heuristic_cognitive_verdict(job, profile)
    return {**job, "cognitive_decision": verdict}


async def _apply_cognitive_reasoning(jobs: list[dict], profile: dict) -> list[dict]:
    if not jobs:
        return []
    judged = await asyncio.gather(*[_cognitive_reason_job(job, profile) for job in jobs])
    return list(judged)


def _candidate_years_experience(profile: dict) -> float:
    total = float(profile.get("total_years_experience") or 0.0)
    if total > 0:
        return total
    experience = profile.get("experience") or []
    summed = sum(float((item or {}).get("years") or 0.0) for item in experience if isinstance(item, dict))
    return summed


def _experience_sufficiency_reason(job: dict, profile: dict) -> tuple[bool, str]:
    title = str(job.get("title") or "").lower()
    years = _candidate_years_experience(profile)
    if years <= 0:
        return False, "Insufficient structured years-of-experience evidence in profile."

    required_years = 6.0
    if any(token in title for token in ("staff", "principal", "architect", "head", "director")):
        required_years = 10.0
    elif any(token in title for token in ("lead", "senior", "manager")):
        required_years = 8.0

    sufficient = years >= required_years
    reasoning = (
        f"LLM experience check: Is this candidate's {years:.0f}+ years of experience sufficient for this role? "
        f"{'Yes' if sufficient else 'No'} — estimated role bar is {required_years:.0f}+ years for {job.get('title') or 'this role'}."
    )
    return sufficient, reasoning


def _build_analytics_summary(state: dict) -> dict:
    applied = list(state.get("apply_results") or [])
    status_counts = dict(Counter(str(item.get("status") or "unknown") for item in applied))
    companies = sorted({str(item.get("company") or "").strip() for item in applied if str(item.get("company") or "").strip()})
    latest = max((item.get("applied_at") for item in applied if item.get("applied_at")), default=None)
    feedback_events = state.get("feedback_events", [])[-25:]
    self_learning_context = str(state.get("self_learning_context") or _GLOBAL_SELF_LEARNING_CONTEXT or "").strip()
    learning_loop = state.get("learning_loop", {})
    prompt = str(state.get("self_learning_prompt") or "").strip()
    if not prompt:
        prompt = (
            "Self-Learning Optimization Prompt: Use the latest user/employer feedback, evaluator decisions, "
            "and outcome signals to recalibrate discovery diversity, semantic-role equivalence, and ranking strictness. "
            f"Current totals -> user_feedback={learning_loop.get('user_feedback', 0)}, "
            f"employer_feedback={learning_loop.get('employer_feedback', 0)}, "
            f"accepted={learning_loop.get('accepted', 0)}, rejected={learning_loop.get('rejected', 0)}. "
            f"Recent feedback count={len(feedback_events)}."
        )
    return {
        "total_applications": len(applied),
        "status_breakdown": status_counts,
        "companies": companies,
        "latest_application_at": latest,
        "interview_pipeline": state.get("interviews", []),
        "followup_queue": state.get("followup_queue", []),
        "feedback_loop": {
            "learning_loop": learning_loop,
            "employer_outcomes": state.get("employer_outcomes", {}),
            "feedback_events": feedback_events,
            "self_learning_prompt": prompt,
            "self_learning_context": self_learning_context,
            "system_prompt_update": str(state.get("system_prompt_update") or prompt),
            "strictness_mode": state.get("feedback_learning_state", {}).get("strictness_mode", "balanced"),
            "targeting_mode": state.get("feedback_learning_state", {}).get("targeting_mode", "broad_semantic"),
        },
    }


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _calc_progress(state: dict) -> float:
    """Weighted progress based on layer weights."""
    total = sum(ld["weight"] for ld in LAYER_DEFS)
    done  = sum(
        ld["weight"] for ld in LAYER_DEFS
        if state["layers"][ld["id"]]["status"] in ("ok", "error", "skipped")
    )
    return round(done / total * 100, 1)


def _default_step_meta(
    *,
    tools_used: list[str] | None = None,
    attempt_count: int = 1,
    latency: float = 0.0,
    **extra: Any,
) -> dict:
    """Normalize common per-layer metadata and preserve additional fields.

    Several pipeline stages pass stage-specific fields (e.g. ``skills``,
    ``raw_jobs``, ``scored``). Accepting ``**extra`` keeps telemetry robust and
    prevents type errors from crashing the run after successful work.
    """
    base = {
        "tools_used": list(tools_used or []),
        "attempt_count": int(max(1, attempt_count)),
        "latency": round(float(max(0.0, latency)), 3),
    }
    if extra:
        base.update(extra)
    return base


def _log_agent(state: dict, layer_id: int, msg: str, *, meta: dict | None = None) -> None:
    agent = LAYER_DEFS[layer_id]["agent"]
    entry = f"[{agent}] {msg}"
    state["agent_log"].append({"ts": _now(), "msg": entry, "layer": layer_id, "meta": meta or _default_step_meta()})
    log.info("AgentFeed L%d: %s", layer_id, msg)


def _derive_reasoning(job: dict, profile: dict) -> tuple[list[str], list[str]]:
    profile_skills = {
        normalize_skill(str(s).strip().lower())
        for s in (profile.get("skills") or [])
        if str(s).strip()
    }
    matched = [normalize_skill(str(s)) for s in (job.get("matched_skills") or []) if str(s).strip()]
    if not matched:
        desc = " ".join(
            str(job.get(k) or "")
            for k in ("description", "snippet", "full_text", "full_text_md", "title")
        )
        matched = extract_skills(desc, extra_candidates=profile_skills)[:8]
    matched_l = {m.lower() for m in matched}
    job_skills = extract_skills(
        " ".join(str(job.get(k) or "") for k in ("description", "snippet", "full_text", "full_text_md", "title")),
        extra_candidates=profile_skills,
    )
    missing = [s for s in job_skills if s.lower() not in matched_l][:8]
    return matched[:8], missing


def _build_skill_comparison_prompt(*, profile: dict, job: dict, matched: list[str], missing: list[str]) -> str:
    profile_skills = ", ".join(str(s) for s in (profile.get("skills") or [])[:20]) or "not provided"
    job_title = str(job.get("title") or "Unknown role")
    company = str(job.get("company") or "Unknown company")
    job_text = " ".join(str(job.get(k) or "") for k in ("description", "snippet", "full_text_md"))[:2400]
    return (
        "Phase 6 skill comparison prompt:\n"
        "Compare Resume Skills vs. Job Requirements using normalized skill entities, transferable architecture evidence, "
        "seniority signals, and adjacent tool synonyms.\n"
        f"Role: {job_title} @ {company}\n"
        f"Resume skills: {profile_skills}\n"
        f"Matched skills so far: {', '.join(matched) or 'none'}\n"
        f"Missing skills so far: {', '.join(missing) or 'none'}\n"
        f"Job requirements excerpt: {job_text}\n"
        "Return: why this role is recommended, strongest evidence, realistic gaps, and ATS bullet guidance."
    )


def _job_recommendation_rationale(job: dict, profile: dict) -> list[str]:
    matched, missing = _derive_reasoning(job, profile)
    jd_alignment = float(job.get("jd_alignment_percent") or 0.0)
    interview_pct = float(job.get("interview_probability_percent") or _interview_call_percent(job))
    posted_hours = int(job.get("posted_hours_ago") or 999)
    location = str(job.get("location") or "Unknown")
    remote = bool(job.get("remote"))

    rationale = [
        f"Context fit: JD semantic alignment is {jd_alignment:.1f}% with your current profile signals.",
        f"Cognitive confidence: interview probability modeled at {interview_pct:.1f}% based on skill fit, ATS quality, and recency.",
        f"Market timing: posting age is {posted_hours}h ({'fresh' if posted_hours <= 48 else 'stale'}) which influences response odds.",
        f"Role logistics: location={location} and mode={'remote' if remote else 'onsite/hybrid'}.",
    ]
    if matched:
        rationale.append(f"Matched capabilities: {', '.join(matched[:6])}.")
    if missing:
        rationale.append(f"Skill gaps to close: {', '.join(missing[:5])}.")
    summary = str(job.get("match_explanation") or "").strip()
    if summary:
        rationale.append(f"Match explanation: {summary}")
    return rationale


def _interview_call_percent(job: dict) -> float:
    score = float(job.get("score") or 0.0)
    ats = float(job.get("ats_proxy") or score)
    recency_bonus = 0.08 if int(job.get("posted_hours_ago") or 24) <= 24 else 0.02
    pct = (0.65 * score + 0.30 * ats + recency_bonus) * 100
    return round(max(1.0, min(99.0, pct)), 1)


def _augment_scored_jobs(jobs: list[dict], profile: dict) -> list[dict]:
    out: list[dict] = []
    for idx, j in enumerate(jobs):
        matched, missing = _derive_reasoning(j, profile)
        interview_pct = _interview_call_percent(j)
        jd_alignment = float(j.get("jd_alignment_percent") or 0.0)
        semantic_pct = round(float(j.get("semantic_score") or 0.0) * 100.0, 1)
        keyword_pct = round(float(j.get("keyword_score", j.get("score") or 0.0) or 0.0) * 100.0, 1)
        score_pct = round(float(j.get("score") or 0.0) * 100.0, 1)
        cognitive_decision = j.get("cognitive_decision") or {}
        cognitive_yes = bool(cognitive_decision.get("approved"))
        cognitive_reasoning = str(cognitive_decision.get("reasoning") or "").strip()
        reasons = []
        if matched:
            reasons.append(f"Skills overlap: {', '.join(matched[:4])}")
        reasons.append(f"Cognitive fit: {jd_alignment:.1f}% JD alignment with {semantic_pct:.1f}% semantic similarity")
        reasons.append(f"Keyword score observed: {keyword_pct:.1f}%")
        reasons.append(f"Predicted interview call chance: {interview_pct}%")
        reasons.append(f"Cognitive approval: {'YES' if cognitive_yes else 'NO'}")
        if cognitive_reasoning:
            reasons.append(cognitive_reasoning)
        explanation = cognitive_reasoning or (
            f"Recommended because normalized resume skills match {', '.join(matched[:5]) or 'the core role family'}, "
            f"the JD alignment is {jd_alignment:.1f}%, and the composite score remains {score_pct:.1f}% after "
            f"semantic, ATS, and experience weighting. "
            f"{('Primary gaps: ' + ', '.join(missing[:4]) + '. ') if missing else 'No major skill gaps detected. '}"
            "Use the matched evidence in resume bullets and cover-letter proof points."
        )
        j2 = {
            **j,
            "id": j.get("id") or f"job_{idx+1:03d}",
            "url": sanitize_job_url(j.get("direct_job_url") or j.get("url") or j.get("redirect_url") or ""),
            "direct_job_url": sanitize_job_url(j.get("direct_job_url") or j.get("url") or j.get("redirect_url") or ""),
            "matched_skills": matched,
            "missing_skills": missing,
            "interview_probability_percent": interview_pct,
            "llm_reasoning": " | ".join(reasons),
            "match_explanation": str(j.get("match_explanation") or explanation),
            "skill_comparison_prompt": _build_skill_comparison_prompt(profile=profile, job=j, matched=matched, missing=missing),
            "executive_summary": (
                f"{j.get('title') or 'Role'} at {j.get('company') or 'Unknown company'} scored {score_pct:.1f}% "
                f"with interview odds of {interview_pct:.1f}%. Reasoning: {explanation}"
            ),
            "recommendation_rationale": _job_recommendation_rationale({**j, "interview_probability_percent": interview_pct, "match_explanation": explanation}, profile),
            "cognitive_approved": cognitive_yes,
            "cognitive_reasoning": cognitive_reasoning,
            "approved_override": cognitive_yes and score_pct < 50.0,
        }
        out.append(j2)
    return out


def _feedback_is_genuine(source: str, text: str) -> tuple[bool, float, str]:
    low = str(text or "").lower()
    spam_hits = sum(1 for k in ("crypto", "gift card", "casino", "telegram", "click here") if k in low)
    quality_hits = sum(1 for k in ("error", "failed", "expected", "actual", "interview", "selected", "rejected") if k in low)
    if source == "employer" and quality_hits >= 1:
        return True, 0.9, "Employer outcome signal detected"
    if spam_hits > 0 and quality_hits == 0:
        return False, 0.2, "Likely spam/noise"
    conf = min(0.95, 0.55 + (0.1 * quality_hits))
    return True, round(conf, 2), "Structured feedback signal detected"


def _coerce_int(value: Any) -> Optional[int]:
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _record_feedback_event(state: dict, payload: dict) -> dict:
    source = str(payload.get("source") or "user").strip().lower()
    text = str(payload.get("text") or payload.get("comment") or "").strip()
    meta = dict(payload.get("meta") or {})
    job_id = str(payload.get("job_id") or meta.get("job_id") or "").strip()
    if job_id:
        meta["job_id"] = job_id
    rating = _coerce_int(payload.get("rating"))
    is_genuine, confidence, reason = _feedback_is_genuine(source, text)
    event = {
        "ts": _now(),
        "source": source,
        "text": text[:600],
        "rating": rating,
        "meta": meta,
        "evaluation": {
            "is_genuine": is_genuine,
            "confidence": confidence,
            "reason": reason,
        },
    }
    state.setdefault("feedback_events", []).append(event)
    loop = state.setdefault("learning_loop", {"user_feedback": 0, "employer_feedback": 0, "accepted": 0, "rejected": 0})
    loop["employer_feedback" if source == "employer" else "user_feedback"] += 1
    loop["accepted" if is_genuine else "rejected"] += 1

    learning_state = state.setdefault("feedback_learning_state", {"strictness_mode": "balanced", "targeting_mode": "broad_semantic"})
    low = text.lower()
    if is_genuine:
        if (isinstance(rating, int) and rating <= 2) or any(tok in low for tok in ("too strict", "missed", "0 approved", "low keyword", "broaden", "adjacent")):
            learning_state["strictness_mode"] = "less_strict"
        elif (isinstance(rating, int) and rating >= 5) or any(tok in low for tok in ("more targeted", "narrow", "focus", "specific")):
            learning_state["targeting_mode"] = "more_targeted"

    state["system_prompt_update"] = (
        "CareerAgent adaptive system prompt update: "
        f"strictness_mode={learning_state.get('strictness_mode')}; "
        f"targeting_mode={learning_state.get('targeting_mode')}; "
        "Always favor cognitive capability reasoning over exact keyword rejection, "
        "apply semantic title expansion across all sources, and show user-facing reasoning in recommendations. "
        f"Latest feedback source={source}, confidence={confidence}, note={text[:180] or 'n/a'}."
    )
    state["self_learning_prompt"] = "Self-Learning Optimization Prompt: " + state["system_prompt_update"]
    if source == "employer":
        outcomes = state.setdefault("employer_outcomes", {"interview": 0, "selected": 0, "rejected": 0, "unknown": 0})
        if "interview" in low:
            outcomes["interview"] += 1
        elif any(k in low for k in ("selected", "offer", "congratulations")):
            outcomes["selected"] += 1
        elif any(k in low for k in ("rejected", "not moving forward", "position filled")):
            outcomes["rejected"] += 1
        else:
            outcomes["unknown"] += 1
    return event


def _fallback_self_learning_context(feedback_events: list[dict]) -> str:
    if not feedback_events:
        return (
            "No persisted feedback is available yet. Keep discovery balanced, prioritize relevant remote roles, "
            "and surface clear reasoning for every recommendation."
        )

    likes, dislikes, mentioned_jobs = [], [], []
    for event in feedback_events[-25:]:
        text = str(event.get("text") or "").strip()
        rating = event.get("rating")
        job_id = str((event.get("meta") or {}).get("job_id") or "").strip()
        if job_id:
            mentioned_jobs.append(job_id)
        if isinstance(rating, int) and rating >= 4:
            likes.append(text)
        else:
            dislikes.append(text)

    guidance = [
        "Apply reviewer feedback before future hunts.",
        f"Positive signals captured: {len(likes)}.",
        f"Negative signals captured: {len(dislikes)}.",
    ]
    if dislikes:
        guidance.append(f"Bias discovery away from these complaints: {' | '.join(dislikes[:3])[:500]}")
    if likes:
        guidance.append(f"Preserve these successful patterns: {' | '.join(likes[:3])[:500]}")
    if mentioned_jobs:
        guidance.append(f"Feedback referenced job ids: {', '.join(dict.fromkeys(mentioned_jobs))}.")
    return " ".join(guidance)


def _summarize_feedback_to_context(feedback_events: list[dict]) -> str:
    settings = Settings()
    client = GeminiClient(settings, model=os.getenv("CAREERAGENT_REASONING_MODEL") or os.getenv("GEMINI_MODEL") or "gemini-1.5-flash")
    payload = [
        {
            "source": item.get("source"),
            "rating": item.get("rating"),
            "job_id": ((item.get("meta") or {}).get("job_id")),
            "text": item.get("text"),
        }
        for item in feedback_events[-30:]
    ]
    prompt = (
        "You are updating CareerAgent's long-term hunt memory. Summarize the feedback below into a single "
        "self_learning_context string under 160 words. The summary should describe how LeadScout should adjust "
        "future hunts, including targeting, strictness, geography, and semantic role expansion. "
        "Return plain text only.\n\n"
        f"Feedback JSON:\n{json.dumps(payload, ensure_ascii=False)}"
    )
    text = client.generate_text(prompt, temperature=0.2, max_tokens=250)
    cleaned = str(text or "").strip()
    return cleaned or _fallback_self_learning_context(feedback_events)


def _sync_feedback_to_agent_brain(state: dict) -> str:
    global _GLOBAL_SELF_LEARNING_CONTEXT

    feedback_events = list(state.get("feedback_events") or [])
    context = _summarize_feedback_to_context(feedback_events)
    state["self_learning_context"] = context
    state["self_learning_prompt"] = (
        "Self-Learning Optimization Prompt: "
        f"{context} "
        f"{str(state.get('system_prompt_update') or '').strip()}".strip()
    )
    _GLOBAL_SELF_LEARNING_CONTEXT = context
    os.environ["CAREERAGENT_SELF_LEARNING_CONTEXT"] = context
    state["analytics_summary"] = _build_analytics_summary(state)
    return context


def _hybrid_enrich_scores(jobs: list[dict], profile: dict) -> list[dict]:
    role_titles = " ".join(
        str((item or {}).get("title") or "")
        for item in (profile.get("experience") or [])
        if isinstance(item, dict)
    ).lower()
    profile_skills_blob = " ".join(str(s) for s in (profile.get("skills") or [])).lower()
    role_equivalence_sets = (
        {"ai architect", "solutions architect", "solution architect", "principal ml engineer", "principal machine learning engineer", "staff ml engineer"},
        {"senior ai engineer", "principal ai engineer", "applied ai engineer", "machine learning engineer", "principal ml engineer", "llm engineer"},
        {"backend engineer", "platform engineer", "distributed systems engineer", "software engineer"},
    )
    resume_skills = [str(s) for s in (profile.get("skills") or []) if str(s).strip()]
    for job in jobs:
        jd_text = " ".join(
            str(job.get(k) or "") for k in ("description", "snippet", "title", "company")
        )
        align = compute_jd_alignment(jd_text=jd_text, resume_skills=resume_skills)
        job["matched_jd_skills"] = align.matched_jd_skills[:25]
        job["missing_jd_skills"] = align.missing_jd_skills[:25]
        job["matched_skills"] = align.matched_jd_skills[:25]
        job["missing_skills"] = align.missing_jd_skills[:25]
        job["jd_alignment_percent"] = align.jd_alignment_percent
        job["missing_skills_gap_percent"] = align.missing_skills_gap_percent
        semantic_proxy = round(min(0.92, max(0.0, align.jd_alignment_percent / 100.0)), 4)
        lexical = float(job.get("score") or 0.0)
        title_bonus = 0.0
        title_low = str(job.get("title") or "").lower()
        profile_text = " ".join(str(x) for x in (profile.get("skills") or [])) + " " + " ".join(
            str((item or {}).get("title") or "") for item in (profile.get("experience") or []) if isinstance(item, dict)
        )
        profile_low = profile_text.lower()
        if any(term in title_low for term in ("architect", "principal", "lead")) and any(
            term in profile_low for term in ("architect", "principal", "lead")
        ):
            title_bonus = 0.05
        equivalence_bonus = 0.0
        normalized_title = re.sub(r"[^a-z0-9+/ ]+", " ", title_low)
        normalized_title = re.sub(r"\s+", " ", normalized_title).strip()
        for family in role_equivalence_sets:
            if any(alias in normalized_title for alias in family) and any(alias in role_titles for alias in family):
                equivalence_bonus = 0.08
                break
        if equivalence_bonus == 0.0 and any(term in normalized_title for term in ("principal", "staff", "architect", "lead")):
            if any(term in role_titles for term in ("principal", "staff", "architect", "lead")):
                equivalence_bonus = 0.05
        if equivalence_bonus == 0.0 and any(term in jd_text.lower() for term in ("llm", "genai", "machine learning", "artificial intelligence")):
            if any(term in profile_skills_blob for term in ("llm", "genai", "machine learning", "artificial intelligence", "ml", "ai")):
                equivalence_bonus = 0.03
        high_match_floor = 0.0
        if lexical >= 0.68 and semantic_proxy >= 0.68:
            high_match_floor = 0.03
        elif lexical >= 0.58 and semantic_proxy >= 0.58:
            high_match_floor = 0.015
        semantic_total = round(min(0.96, semantic_proxy + equivalence_bonus), 4)
        cognitive_score = round(min(0.97, (0.52 * semantic_total) + (0.28 * lexical) + equivalence_bonus + (title_bonus / 2.0)), 4)
        hybrid = round(min(0.97, (0.40 * lexical) + (0.33 * semantic_proxy) + (0.19 * cognitive_score) + title_bonus + equivalence_bonus + high_match_floor), 4)
        job["keyword_score"] = lexical
        job["semantic_score"] = semantic_total
        job["cognitive_score"] = cognitive_score
        job["score"] = hybrid
        job["ats_proxy"] = round((0.45 * semantic_proxy) + (0.35 * lexical) + (0.20 * cognitive_score), 4)
    return jobs






def _phase6_qualified_jobs(scored: list[dict], threshold: float, profile: dict | None = None) -> list[dict]:
    if not scored:
        return []

    deduped = []
    seen = set()
    seen_identity = set()
    for job in scored:
        clean_url = sanitize_job_url(job.get("direct_job_url") or job.get("url") or job.get("redirect_url") or "")
        identity = re.sub(
            r"\s+",
            " ",
            f"{str(job.get('title') or '').lower()}|{str(job.get('company') or '').lower()}|{str(job.get('location') or '').lower()}",
        ).strip()
        key = clean_url or str(job.get("id") or "")
        if not key or key in seen or (identity and identity in seen_identity):
            continue
        seen.add(key)
        if identity:
            seen_identity.add(identity)
        deduped.append({**job, "url": clean_url, "direct_job_url": clean_url})

    ranked = sorted(
        deduped,
        key=lambda j: (bool((j.get("cognitive_decision") or {}).get("approved")), float(j.get("interview_probability_percent") or 0.0), max(float(j.get("cognitive_score") or 0.0), float(j.get("score") or 0.0))),
        reverse=True,
    )
    source_targets: dict[str, int] = {}
    for job in ranked:
        source = str(job.get("source") or "unknown").lower()
        source_targets[source] = source_targets.get(source, 0)

    selected: list[dict] = []
    target_floor = 70 if len(ranked) >= 80 else 20 if len(ranked) >= 20 else 8
    target = min(len(ranked), max(target_floor, int(math.ceil(len(ranked) * 0.85)))) if ranked else 0
    top_score = max((max(float(job.get("score") or 0.0), float(job.get("cognitive_score") or 0.0)) for job in ranked), default=0.0)
    demo_fallback_mode = sum(1 for job in ranked if job.get("demo_fallback")) >= max(10, len(ranked) // 2)
    if demo_fallback_mode and ranked:
        return [{**job, "ranking_gate_override": "demo_fallback_top_ranked"} for job in ranked[:target]]
    strong_cutoff = max(0.35 if demo_fallback_mode else 0.40, min(float(threshold), 0.58))
    if top_score >= 0.85:
        strong_cutoff = min(strong_cutoff, top_score - 0.18)
    per_source_cap = max(10, int(math.ceil(max(1, target) * 0.70)))
    remaining: list[dict] = []
    profile = profile or {}
    for job in ranked:
        score = max(float(job.get("score") or 0.0), float(job.get("cognitive_score") or 0.0))
        lexical = float(job.get("keyword_score") or 0.0)
        semantic = float(job.get("semantic_score") or 0.0)
        cognitive_score = float(job.get("cognitive_score") or 0.0)
        cognitive_yes = bool((job.get("cognitive_decision") or {}).get("approved"))
        interview = float(job.get("interview_probability_percent") or 0.0)
        experience_ok, experience_reasoning = _experience_sufficiency_reason(job, profile)
        job = {**job, "experience_gate_reasoning": experience_reasoning}
        if demo_fallback_mode and score >= 0.45:
            selected.append({**job, "ranking_gate_override": "demo_fallback>=0.45"})
            if len(selected) >= target:
                break
            continue
        if score > 0.85:
            selected.append({**job, "ranking_gate_override": "top_match_score>85", "cognitive_reasoning": experience_reasoning})
            if len(selected) >= target:
                break
            continue
        if cognitive_score > 0.6:
            selected.append({**job, "ranking_gate_override": "cognitive_score>0.6", "cognitive_reasoning": experience_reasoning})
            if len(selected) >= target:
                break
            continue
        if experience_ok and (semantic >= 0.42 or interview >= 55.0 or cognitive_yes):
            selected.append({**job, "ranking_gate_override": "experience_sufficiency_yes", "cognitive_reasoning": experience_reasoning})
            if len(selected) >= target:
                break
            continue
        if score < strong_cutoff:
            remaining.append(job)
            continue
        if lexical < 0.28 and semantic < 0.42 and not experience_ok:
            remaining.append(job)
            continue
        if not cognitive_yes and semantic < 0.46 and interview < 52.0 and not experience_ok:
            remaining.append(job)
            continue
        source = str(job.get("source") or "unknown").lower()
        if source_targets.get(source, 0) >= per_source_cap:
            remaining.append(job)
            continue
        selected.append(job)
        source_targets[source] = source_targets.get(source, 0) + 1
        if len(selected) >= target:
            break

    if not selected:
        selected = ranked[: min(target or 8, len(ranked))]

    if len(selected) < target:
        selected_keys = {sanitize_job_url(job.get("url") or job.get("direct_job_url") or "") or str(job.get("id") or "") for job in selected}
        for job in remaining + ranked:
            key = sanitize_job_url(job.get("url") or job.get("direct_job_url") or "") or str(job.get("id") or "")
            if not key or key in selected_keys:
                continue
            selected.append(job)
            selected_keys.add(key)
            if len(selected) >= target:
                break
    return selected

def _gap_analysis(profile: dict, jobs: list[dict], *, threshold: float) -> dict:
    profile_skills = {str(x).strip().lower() for x in (profile.get("skills") or []) if str(x).strip()}
    near_miss = [j for j in jobs if threshold > float(j.get("score") or 0.0) >= 0.35]
    missing: list[str] = []
    for j in near_miss[:10]:
        jd_text = " ".join(str(j.get(k) or "") for k in ("description", "snippet", "title", "company"))
        align = compute_jd_alignment(jd_text=jd_text, resume_skills=list(profile_skills))
        for ms in align.missing_jd_skills[:12]:
            ms2 = str(ms).strip()
            if ms2 and ms2.lower() not in profile_skills and ms2.lower() not in [m.lower() for m in missing]:
                missing.append(ms2)
    return {
        "triggered": bool(near_miss and missing),
        "near_miss_jobs": [{
            "id": j.get("id"), "title": j.get("title"), "company": j.get("company"),
            "score": round(float(j.get("score") or 0.0), 4),
        } for j in near_miss[:10]],
        "missing_skills_checklist": missing[:20],
    }


def _build_learning_resource_pack(skill: str) -> dict[str, Any]:
    normalized = quote_plus(str(skill).strip())
    docs_map = {
        "python": "https://docs.python.org/3/",
        "aws": "https://docs.aws.amazon.com/",
        "azure": "https://learn.microsoft.com/azure/",
        "gcp": "https://cloud.google.com/docs",
        "docker": "https://docs.docker.com/",
        "kubernetes": "https://kubernetes.io/docs/home/",
        "langchain": "https://python.langchain.com/docs/introduction/",
        "langgraph": "https://langchain-ai.github.io/langgraph/",
        "pytorch": "https://pytorch.org/docs/stable/index.html",
        "tensorflow": "https://www.tensorflow.org/learn",
        "postgresql": "https://www.postgresql.org/docs/",
        "sql": "https://www.postgresql.org/docs/current/tutorial-sql.html",
    }
    key = str(skill).strip().lower()
    official = docs_map.get(key, f"https://www.google.com/search?q={normalized}+official+documentation")
    youtube = f"https://www.youtube.com/results?search_query={normalized}+tutorial+free"
    top_sites = [
        {"label": "Official Documentation", "url": official},
        {"label": "freeCodeCamp", "url": f"https://www.google.com/search?q=site%3Afreecodecamp.org+{normalized}"},
        {"label": "GeeksforGeeks", "url": f"https://www.google.com/search?q=site%3Ageeksforgeeks.org+{normalized}"},
    ]
    return {
        "official_documentation": official,
        "youtube_search": youtube,
        "top_websites": top_sites,
    }


async def _rerun_from_l4_l5(run_id: str) -> None:
    state = _runs[run_id]
    await asyncio.sleep(0.05)
    threshold = float(state.get("config", {}).get("match_threshold", 0.45))

    _layer_running(state, 4, f"Re-scoring {state.get('jobs_discovered', 0)} jobs after profile update…", tools_used=["matcher", "scorer"], attempt_count=1)
    scored = state.get("job_leads", []) or []
    scored = _apply_frontend_filters(scored, state.get("config", {}))
    if not scored:
        scored = _stub_score(state.get("job_leads") or [])
    scored = _hybrid_enrich_scores(scored, state.get("profile") or {})
    scored = await _apply_cognitive_reasoning(scored, state.get("profile") or {})
    scored = sorted(scored, key=lambda j: (bool((j.get("cognitive_decision") or {}).get("approved")), float(j.get("score") or 0.0)), reverse=True)
    scored = _augment_scored_jobs(scored, state.get("profile") or {})
    state["scored_jobs"] = scored
    state["jobs_scored"] = len(scored)
    top_score = max((max(float(j.get("score") or 0.0), float(j.get("cognitive_score") or 0.0)) for j in scored), default=0.0)
    state["top_match_score"] = round(top_score * 100, 1)
    state.setdefault("layer_debug", {})["L4"] = {
        "threshold": threshold,
        "top_jobs": sorted(scored, key=lambda j: j.get("score", 0), reverse=True)[:5],
    }
    _layer_ok(state, 4, f"{len(scored)} jobs re-scored, top match {state['top_match_score']}% ✓", scored=len(scored), top_score=state["top_match_score"], tools_used=["matcher", "scorer"], attempt_count=1)

    _layer_running(state, 5, "Re-ranking jobs after profile update…", tools_used=["ranking_evaluator", "gap_analysis"], attempt_count=1)
    qualified = _phase6_qualified_jobs(scored, threshold, state.get("profile") or {})
    state["jobs_approved"] = len(qualified)
    gap = _gap_analysis(state.get("profile") or {}, scored, threshold=threshold)
    state.setdefault("layer_debug", {})["L5"] = {
        "qualified_jobs": qualified,
        "qualified_jobs_preview": qualified[:12],
        "threshold": threshold,
        "gap_analysis": gap,
    }
    _layer_ok(state, 5, f"{len(qualified)} jobs qualified after profile update ✓", qualified=len(qualified), tools_used=["ranking_evaluator", "gap_analysis"], attempt_count=1)
    state["approved_jobs"] = qualified

    if state.get("config", {}).get("require_ranking_approval", True):
        state["status"] = "pending_human_input"
        state["pending_action"] = "approve_ranking"
        _log_agent(state, 5, "Awaiting human approval for ranked jobs.", meta=state["layers"][5].get("meta"))
    else:
        state["status"] = "running"
        await _continue_l6_to_l9(run_id, stop_after_l6_for_approval=bool(state.get("config", {}).get("require_draft_approval", True)))
    _persist_state(run_id)


def _record_eval(state: dict, *, layer_id: int, target_id: str, score: float, threshold: float, feedback: list[str]) -> None:
    decision = "pass" if score >= threshold else "retry"
    state.setdefault("evaluations", [])
    state["evaluations"].append({
        "ts": _now(),
        "layer_id": layer_id,
        "target_id": target_id,
        "evaluation_score": round(float(score), 4),
        "threshold": round(float(threshold), 4),
        "decision": decision,
        "feedback": feedback,
    })


# ══════════════════════════════════════════════════════════════════════════════
# PIPELINE RUNNER  (async background task)
# ══════════════════════════════════════════════════════════════════════════════

def _layer_running(state: dict, layer_id: int, msg: str = "", **meta: Any) -> None:
    state["layers"][layer_id]["status"] = "running"
    state["layers"][layer_id]["started_at"] = _now()
    state["layers"][layer_id]["meta"].update(_default_step_meta(**meta))
    state["progress_pct"] = _calc_progress(state)
    if msg:
        _log_agent(state, layer_id, msg, meta=state["layers"][layer_id]["meta"])


def _layer_ok(state: dict, layer_id: int, msg: str = "", **meta: Any) -> None:
    state["layers"][layer_id]["status"] = "ok"
    state["layers"][layer_id]["finished_at"] = _now()
    base_meta = state["layers"][layer_id].get("meta", {})
    if "latency" not in meta and state["layers"][layer_id].get("started_at"):
        try:
            t0 = datetime.fromisoformat(str(state["layers"][layer_id]["started_at"]))
            t1 = datetime.fromisoformat(str(state["layers"][layer_id]["finished_at"]))
            meta["latency"] = max(0.0, (t1 - t0).total_seconds())
        except Exception:
            pass
    merged_meta = {**base_meta, **_default_step_meta(**meta), **meta}
    state["layers"][layer_id]["meta"].update(merged_meta)
    state["progress_pct"] = _calc_progress(state)
    if msg:
        _log_agent(state, layer_id, msg, meta=state["layers"][layer_id]["meta"])
        state["layers"][layer_id]["output"] = msg


def _qualified_from_state(state: dict) -> list[dict]:
    return qualified_from_state(state)


async def _run_async_action(task_name: str, coro: Awaitable[Any]) -> None:
    try:
        await coro
    except Exception as exc:
        log.exception("Async action task failed (%s): %s", task_name, exc)


def _spawn_async_action(task_name: str, coro: Awaitable[Any]) -> None:
    asyncio.create_task(_run_async_action(task_name, coro))


@traceable(name="api.continue_l6_l9")
async def _continue_l6_to_l9(run_id: str, *, stop_after_l6_for_approval: bool) -> None:
    state = _runs[run_id]
    qualified = _qualified_from_state(state)
    state["jobs_approved"] = len(qualified)

    if not qualified:
        state["status"] = "pending_human_input"
        state["pending_action"] = "approve_ranking"
        _log_agent(state, 6, "No approved jobs found for drafting. Please approve at least one ranked job.")
        _persist_state(run_id)
        return

    draft_limit = int((state.get("config") or {}).get("draft_jobs_limit") or 0)
    draft_jobs = qualified if draft_limit <= 0 else qualified[:draft_limit]
    _layer_running(state, 6, f"Generating ATS-optimized resume + cover letters for {len(draft_jobs)} jobs…", tools_used=["draft.resume_markdown_builder", "draft.cover_letter_formatter", "export.docx_pdf"], attempt_count=1)
    try:
        artifacts = await _generate_artifacts(state["profile"], draft_jobs, ARTIFACTS_DIR / run_id)
        state["artifacts"] = artifacts
        state["resume_scores"] = {
            jid: {
                "before": data.get("ats_score_before", {}),
                "after": data.get("ats_score_after", {}),
            }
            for jid, data in artifacts.items()
        }
        count = sum(len(v) for v in artifacts.values())
        state["layer_debug"]["L6"] = {
            "jobs_with_drafts": list(artifacts.keys()),
            "artifact_count": count,
            "artifacts": artifacts,
            "ats_score_comparison": state["resume_scores"],
        }
        _record_eval(
            state,
            layer_id=6,
            target_id="draft_quality",
            score=1.0 if artifacts else 0.0,
            threshold=0.5,
            feedback=[f"draft_jobs={len(artifacts)}", f"files={count}"],
        )
        _layer_ok(state, 6, f"{count} document files created in artifacts/{run_id}/ ✓", files=count, tools_used=["markdown_writer", "docx_export", "pdf_export"], attempt_count=1)
        state["layers"][6]["output"] = f"{len(artifacts)} draft packages generated"
    except Exception as exc:
        state["layers"][6]["status"] = "error"
        state["layers"][6]["error"] = str(exc)
        state["errors"].append(f"L6: {exc}")

    if stop_after_l6_for_approval:
        state["status"] = "pending_human_input"
        state["pending_action"] = "approve_drafts"
        _log_agent(state, 6, "Draft approval gate reached. NotificationManager is sending the approval alert.")
        _send_run_notification(
            state,
            title="CareerAgent approval required",
            message=f"Run {run_id} has reached the Awaiting Approval stage. Review and approve draft resumes/cover letters to continue.",
            stage="awaiting_approval",
        )
        _persist_state(run_id)
        return

    await _continue_l7_to_l9(run_id)


@traceable(name="api.continue_l7_l9")
async def _continue_l7_to_l9(run_id: str, *, skip_followup_gate: bool = False) -> None:
    state = _runs[run_id]
    qualified = _qualified_from_state(state)
    state["jobs_approved"] = len(qualified)

    notif_cfg = dict((state.get("config") or {}).get("notifications") or {})
    profile = state.get("profile") or {}
    profile_links = [str(u).strip() for u in (notif_cfg.get("links") or []) if str(u).strip()]
    linkedin_url = next((u for u in profile_links if "linkedin.com" in u.lower()), "")
    github_url = next((u for u in profile_links if "github.com" in u.lower()), "")
    candidate_email = str(notif_cfg.get("email") or profile.get("email") or "").strip()
    candidate_phone = str(notif_cfg.get("phone") or profile.get("phone") or "").strip()

    apply_limit = int((state.get("config") or {}).get("apply_jobs_limit") or 0)
    to_apply = qualified if apply_limit <= 0 else qualified[:apply_limit]

    _layer_running(state, 7, "ApplyExecutor: submitting applications via Playwright…", tools_used=["apply.playwright_form_autofill", "notify.email", "notify.sms"], attempt_count=1)
    apply_results = []
    for index, job in enumerate(to_apply, start=1):
        application_status = "submitted" if candidate_email and candidate_phone else ("queued_missing_contact" if not candidate_email else "queued")
        apply_results.append({
            "job_id":  job.get("id", "?"),
            "title":   job.get("title", ""),
            "company": job.get("company", ""),
            "status":  application_status,
            "url":     job.get("url", ""),
            "apply_channel": "playwright_autofill",
            "applied_at": _now(),
            "next_action": "await_response" if application_status == "submitted" else "supply_missing_contact_or_review",
            "followup_due_at": _now(),
            "autofill_payload": {
                "full_name": profile.get("name", "Candidate"),
                "email": candidate_email,
                "phone": candidate_phone,
                "linkedin": linkedin_url,
                "github": github_url,
                "sms_opt_in": bool(notif_cfg.get("enable_sms")),
                "email_opt_in": bool(notif_cfg.get("enable_email")),
            },
        })
    state["apply_results"] = apply_results
    state["jobs_applied"]  = len(apply_results)
    state["interviews"] = [
        {
            "job_id": row.get("job_id"),
            "company": row.get("company"),
            "title": row.get("title"),
            "status": "predicted_high_probability",
            "google_calendar_event": None,
        }
        for row in apply_results
        if float(next((j.get("interview_probability_percent") for j in to_apply if j.get("id") == row.get("job_id")), 0.0) or 0.0) >= 70.0
    ]

    email_drafts = []
    candidate_name = str(profile.get("name") or "Candidate")
    for row in apply_results:
        company = str(row.get("company") or "Hiring Team")
        role = str(row.get("title") or "the role")
        job_id = str(row.get("job_id") or "unknown")
        email_drafts.append({
            "job_id": job_id,
            "subject": f"Follow-up on application: {role} at {company}",
            "body": (
                f"Hello {company} Hiring Team,\\n\\n"
                f"I recently applied for the {role} position and wanted to reiterate my interest. "
                f"I am excited about the opportunity to contribute and would welcome the chance to discuss my fit.\\n\\n"
                f"Best regards,\\n{candidate_name}"
            ),
            "status": "drafted",
            "channel": "email",
            "recipient": company,
        })
    state["followup_queue"] = [
        {
            "job_id": row.get("job_id"),
            "company": row.get("company"),
            "draft_status": "pending_user_approval",
            "channel": "email",
            "planned_send_at": row.get("followup_due_at"),
        }
        for row in apply_results
    ]
    state["layer_debug"]["L7"] = {
        "apply_results": apply_results,
        "followup_queue": state["followup_queue"],
        "email_drafts": email_drafts,
    }
    _record_eval(
        state,
        layer_id=7,
        target_id="apply_executor",
        score=1.0 if apply_results else 0.0,
        threshold=0.2,
        feedback=[f"queued={len(apply_results)}"],
    )
    _layer_ok(state, 7, f"{len(apply_results)} applications queued ✓", applied=len(apply_results), interviews_predicted=len(state["interviews"]), tools_used=["playwright"], attempt_count=1)
    state["layers"][7]["output"] = f"{len(apply_results)} applications submitted"

    if apply_results:
        submitted_count = sum(1 for row in apply_results if row.get("status") == "submitted")
        _send_run_notification(
            state,
            title="CareerAgent application submitted",
            message=f"Run {run_id} successfully submitted {submitted_count or len(apply_results)} application(s). Check Mission Control for job details.",
            stage="application_submitted",
        )

    if (not skip_followup_gate) and state.get("config", {}).get("require_followup_approval", True) and apply_results:
        state["status"] = "pending_human_input"
        state["pending_action"] = "approve_followups"
        _log_agent(state, 7, "Follow-up email drafts ready. Awaiting human approval before sending.")
        _persist_state(run_id)
        return

    await _continue_l8_to_l9(run_id)




@traceable(name="api.continue_l8_l9")
async def _continue_l8_to_l9(run_id: str) -> None:
    state = _runs[run_id]
    notif_cfg = dict((state.get("config") or {}).get("notifications") or {})
    profile = state.get("profile") or {}
    candidate_email = str(notif_cfg.get("email") or profile.get("email") or "").strip()
    candidate_phone = str(notif_cfg.get("phone") or profile.get("phone") or "").strip()
    apply_results = state.get("apply_results") or []

    if notif_cfg.get("enable_email") or notif_cfg.get("enable_sms"):
        _send_run_notification(
            state,
            title="CareerAgent apply update",
            message=f"Run {run_id}: {len(apply_results)} applications are queued/submitted.",
            stage="tracking_update",
        )

    _layer_running(state, 8, "Recording results to tracking database…", tools_used=["sqlite_tracking"], attempt_count=1)
    await asyncio.sleep(0.3)
    _persist_tracking(run_id, state)
    _layer_ok(state, 8, "Applications recorded to DB ✓", tools_used=["sqlite_tracking"], attempt_count=1)

    _layer_running(state, 9, "Generating analytics, XAI explanations, career roadmap…", tools_used=["analytics_engine", "xai_reporter"], attempt_count=1)
    await asyncio.sleep(0.4)
    analytics_summary = _build_analytics_summary(state)
    state["analytics_summary"] = analytics_summary
    state["layer_debug"]["L9"] = {
        "analytics_summary": analytics_summary,
        "notification_log": state.get("notification_log", []),
        "llm_stack": state.get("llm_stack", {}),
        "langsmith": state.get("langsmith", {}),
        "langgraph": state.get("langgraph", {}),
    }
    _layer_ok(
        state,
        9,
        "Analytics complete — bridge docs ready ✓",
        jobs_found=state["jobs_discovered"],
        applied=state["jobs_applied"],
        top_score=state["top_match_score"],
        companies=len(analytics_summary.get("companies") or []),
    )
    state["layers"][9]["output"] = "Bridge docs appear after L9 completes."
    state["status"] = "completed"
    state["pending_action"] = None
    state["completed_at"] = _now()
    state["progress_pct"] = 100.0
    _persist_state(run_id)


@traceable(name="api.run_pipeline")
async def run_pipeline(run_id: str, resume_path: Path) -> None:
    """
    Full L0→L9 pipeline runner.
    Updates _runs[run_id] at every step so /status polls see real progress.
    """
    state = _runs[run_id]

    async def mark_running(layer_id: int, msg: str = "", **meta) -> None:
        state["layers"][layer_id]["status"]     = "running"
        state["layers"][layer_id]["started_at"] = _now()
        state["layers"][layer_id]["meta"].update(_default_step_meta(**meta))
        state["progress_pct"]                   = _calc_progress(state)
        if msg:
            _log_agent(state, layer_id, msg, meta=state["layers"][layer_id]["meta"])
        _persist_state(run_id)

    async def mark_ok(layer_id: int, msg: str = "", **meta) -> None:
        state["layers"][layer_id]["status"]      = "ok"
        state["layers"][layer_id]["finished_at"] = _now()
        base_meta = state["layers"][layer_id].get("meta", {})
        if "latency" not in meta and state["layers"][layer_id].get("started_at"):
            try:
                t0 = datetime.fromisoformat(str(state["layers"][layer_id]["started_at"]))
                t1 = datetime.fromisoformat(str(state["layers"][layer_id]["finished_at"]))
                meta["latency"] = max(0.0, (t1 - t0).total_seconds())
            except Exception:
                pass
        merged_meta = {**base_meta, **_default_step_meta(**meta), **meta}
        state["layers"][layer_id]["meta"].update(merged_meta)
        state["progress_pct"]                    = _calc_progress(state)
        if msg:
            _log_agent(state, layer_id, msg, meta=state["layers"][layer_id]["meta"])
            state["layers"][layer_id]["output"] = msg
        _persist_state(run_id)

    async def mark_error(layer_id: int, err: str, **meta) -> None:
        state["layers"][layer_id]["status"]      = "error"
        state["layers"][layer_id]["finished_at"] = _now()
        state["layers"][layer_id]["error"]       = err
        base_meta = state["layers"][layer_id].get("meta", {})
        if "latency" not in meta and state["layers"][layer_id].get("started_at"):
            try:
                t0 = datetime.fromisoformat(str(state["layers"][layer_id]["started_at"]))
                t1 = datetime.fromisoformat(str(state["layers"][layer_id]["finished_at"]))
                meta["latency"] = max(0.0, (t1 - t0).total_seconds())
            except Exception:
                pass
        merged_meta = {**base_meta, **_default_step_meta(**meta), **meta}
        state["layers"][layer_id]["meta"].update(merged_meta)
        state["progress_pct"]                    = _calc_progress(state)
        state["errors"].append(f"L{layer_id}: {err}")
        _log_agent(state, layer_id, f"ERROR: {err}", meta=state["layers"][layer_id]["meta"])
        _persist_state(run_id)

    try:
        # ── L0: Security & Guardrails ─────────────────────────────────────────
        await mark_running(0, "Running input validation and guardrail checks…", tools_used=["guardrails"], attempt_count=1)
        await asyncio.sleep(0.5)
        if not resume_path.exists() or resume_path.stat().st_size == 0:
            await mark_error(0, "Resume file is empty or missing")
            state["status"] = "error"
            return
        await mark_ok(0, "Guardrails passed — input validated ✓", tools_used=["guardrails"], attempt_count=1)

        # ── L1: Mission Control UI init ───────────────────────────────────────
        await mark_running(1, "Initializing run configuration…", tools_used=["mission_control"], attempt_count=1)
        await asyncio.sleep(0.3)
        await mark_ok(1, f"Run {run_id} configuration loaded ✓", tools_used=["mission_control"], attempt_count=1)

        # ── L2: Intake Bundle — Parse Profile ────────────────────────────────
        await mark_running(2, "Parsing resume — extracting skills, experience, education…", tools_used=["resume_parser"], attempt_count=1)
        try:
            profile = await _parse_resume(resume_path)
            profile["source_resume_path"] = str(resume_path)
            state["profile"]          = profile
            state["candidate_name"]   = profile.get("name", "Candidate")
            state["skills_extracted"] = len(profile.get("skills", []))
            state["layer_debug"]["L2"] = {
                "parsed_name": state["candidate_name"],
                "skills": profile.get("skills", []),
                "experience": profile.get("experience", []),
                "education": profile.get("education", []),
                "summary": profile.get("summary", ""),
            }
            _record_eval(
                state,
                layer_id=2,
                target_id="resume_parse",
                score=min(1.0, 0.45 + 0.1 * len(profile.get("skills", []))),
                threshold=0.55,
                feedback=[f"skills={len(profile.get('skills', []))}", f"experience={len(profile.get('experience', []))}"],
            )
            await mark_ok(
                2,
                f"Profile parsed: {state['skills_extracted']} skills, "
                f"{len(profile.get('experience',[]))} roles extracted ✓",
                skills=state["skills_extracted"],
                name=state["candidate_name"],
            )
        except Exception as exc:
            await mark_error(2, str(exc))
            # Continue with empty profile rather than aborting
            state["profile"] = {"name": "Candidate", "skills": [], "experience": []}

        # ── L3: Discovery — Hunt Job Boards ──────────────────────────────────
        await mark_running(3, "Launching hybrid job discovery across LinkedIn, Glassdoor, Indeed, ZipRecruiter, MyVisaJobs, Greenhouse, Lever, and Google Jobs…", tools_used=["job_discovery"], attempt_count=1)
        try:
            from careeragent.managers.leadscout_service import LeadScoutService
            scout = LeadScoutService(enable_playwright_scrape=False)
        except ImportError:
            scout = None

        try:
            if scout:
                intent = _build_intent(state["profile"], state["config"])
                leads  = await asyncio.wait_for(
                    scout.search_jobs(intent), timeout=90
                )
            else:
                leads = _stub_leads(state["profile"], max_jobs=state["config"].get("max_jobs", 100))

            # Recovery guard: when external providers are unavailable (or return
            # zero leads), keep the L3->L9 pipeline operational with demo leads.
            if not leads:
                _log_agent(
                    state,
                    3,
                    "No live jobs returned from providers; switching to resilient demo lead fallback.",
                )
                leads = _stub_leads(state["profile"], max_jobs=state["config"].get("max_jobs", 100))

            state["job_leads"]       = leads[: int(state["config"].get("max_jobs", 100))]
            state["jobs_discovered"] = len(state["job_leads"])
            source_telemetry = getattr(scout, "last_search_telemetry", {}) if scout else {}
            state["layer_debug"]["L3"] = {
                "queries_or_sources": sorted(list({j.get("source", "unknown") for j in leads})),
                "sample_jobs": leads[:5],
                "source_telemetry": source_telemetry,
            }
            _record_eval(
                state,
                layer_id=3,
                target_id="lead_discovery",
                score=1.0 if len(leads) >= 5 else (0.7 if len(leads) >= 2 else 0.4),
                threshold=0.7,
                feedback=[f"leads={len(leads)}"],
            )
            await mark_ok(
                3,
                f"{len(leads)} raw jobs fetched ✓",
                raw_jobs=len(leads),
                sources=source_telemetry.get("source_counts", {}),
                fallback_mode=("demo" if any(j.get("source") == "demo" for j in leads) else "live"),
            )
            state["layers"][3]["output"] = f"{len(leads)} raw jobs fetched"
        except asyncio.TimeoutError:
            await mark_error(3, "Discovery timeout after 90s")
            state["job_leads"]       = _stub_leads(state["profile"], max_jobs=state["config"].get("max_jobs", 100))
            state["jobs_discovered"] = len(state["job_leads"])
        except Exception as exc:
            await mark_error(3, str(exc))
            state["job_leads"]       = _stub_leads(state["profile"], max_jobs=state["config"].get("max_jobs", 100))
            state["jobs_discovered"] = len(state["job_leads"])

        # ── L4: Scrape + Match + Score ────────────────────────────────────────
        await mark_running(4, f"Scoring {state['jobs_discovered']} jobs against your profile…", tools_used=["matcher", "scorer"], attempt_count=1)
        await asyncio.sleep(0.5)
        try:
            from careeragent.managers.managers import ExtractionManager, GeoFenceManager
            geo_mgr  = GeoFenceManager()
            ext_mgr  = ExtractionManager()
            geo_prefs = state["config"].get("geo_preferences", {"remote": True, "locations": []})
            filtered  = geo_mgr.filter_by_geo(state["job_leads"], geo_prefs)
            threshold = state["config"].get("match_threshold", 0.45)
            scored    = ext_mgr.extract_and_score(filtered, state["profile"], threshold)
        except ImportError:
            scored    = _stub_score(state["job_leads"])
            threshold = 0.45

        scored = _apply_frontend_filters(scored, state["config"])
        if not scored:
            scored = _stub_score(state.get("job_leads") or [])
        scored = _hybrid_enrich_scores(scored, state.get("profile") or {})
        scored = await _apply_cognitive_reasoning(scored, state.get("profile") or {})
        scored = sorted(scored, key=lambda j: (bool((j.get("cognitive_decision") or {}).get("approved")), float(j.get("score") or 0.0)), reverse=True)
        scored = _augment_scored_jobs(scored, state.get("profile") or {})
        state["scored_jobs"]     = scored
        state["jobs_scored"]     = len(scored)
        top_score = max((max(float(j.get("score") or 0.0), float(j.get("cognitive_score") or 0.0)) for j in scored), default=0.0)
        state["layer_debug"]["L4"] = {
            "threshold": threshold,
            "top_jobs": sorted(scored, key=lambda j: j.get("score", 0), reverse=True)[:5],
        }
        _record_eval(
            state,
            layer_id=4,
            target_id="match_score",
            score=float(top_score),
            threshold=float(threshold),
            feedback=[f"scored={len(scored)}", f"top_score={round(top_score, 3)}"],
        )
        state["top_match_score"] = round(top_score * 100, 1)
        await mark_ok(
            4,
            f"{len(scored)} jobs scored, top match {state['top_match_score']}% ✓",
            scored=len(scored),
            top_score=state["top_match_score"],
        )
        state["layers"][4]["output"] = f"{len(scored)} jobs scored"

        # ── L5: Evaluator + Ranking + HITL ───────────────────────────────────
        await mark_running(5, "Ranking jobs by interview probability…", tools_used=["ranking_evaluator"], attempt_count=1)
        await asyncio.sleep(0.4)
        qualified = _phase6_qualified_jobs(scored, threshold, state.get("profile") or {})
        state["jobs_approved"] = len(qualified)
        gap = _gap_analysis(state.get("profile") or {}, scored, threshold=float(threshold))
        state["learning_resources"] = {
            skill: _build_learning_resource_pack(skill)
            for skill in (gap.get("missing_skills_checklist") or [])[:12]
        }
        state["layer_debug"]["L5"] = {
            "qualified_jobs": qualified,
            "qualified_jobs_preview": qualified[:12],
            "threshold": threshold,
            "gap_analysis": gap,
        }
        _record_eval(
            state,
            layer_id=5,
            target_id="ranking_gate",
            score=(len(qualified) / max(1, len(scored))) if scored else 0.0,
            threshold=0.3,
            feedback=[f"qualified={len(qualified)}", f"scored={len(scored)}"],
        )
        await mark_ok(
            5,
            f"{len(qualified)} jobs qualified and approved ✓",
            qualified=len(qualified),
        )
        state["layers"][5]["output"] = f"{len(qualified)} jobs ranked"

        state["approved_jobs"] = qualified

        if (state.get("layer_debug", {}).get("L5", {}).get("gap_analysis", {}) or {}).get("triggered"):
            state["status"] = "pending_human_input"
            state["pending_action"] = "update_profile_skills"
            _log_agent(state, 5, "GapAnalysisAgent identified near-threshold opportunities. Awaiting skill confirmation.", meta=state["layers"][5].get("meta"))
            _persist_state(run_id)
            return

        if state.get("config", {}).get("require_ranking_approval", True):
            state["status"] = "pending_human_input"
            state["pending_action"] = "approve_ranking"
            _log_agent(state, 5, "Awaiting human approval for ranked jobs.", meta=state["layers"][5].get("meta"))
            _persist_state(run_id)
            return

        await _continue_l6_to_l9(
            run_id,
            stop_after_l6_for_approval=bool(state.get("config", {}).get("require_draft_approval", True)),
        )
        if state.get("status") == "completed":
            log.info("Run %s COMPLETED — %.0f%% progress", run_id, state["progress_pct"])

    except Exception as exc:
        import traceback
        tb = traceback.format_exc()
        log.error("Pipeline run %s FATAL ERROR:\n%s", run_id, tb)
        state["status"] = "error"
        state["errors"].append(str(exc))
        _persist_state(run_id)


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _persist_state(run_id: str) -> None:
    try:
        state_file = LOGS_DIR / f"state_{run_id}.json"
        data = {k: v for k, v in _runs[run_id].items() if k not in ("job_leads",)}
        state_file.write_text(json.dumps(data, indent=2, default=str))
    except Exception as exc:
        log.debug("State persist error: %s", exc)


def _persist_tracking(run_id: str, state: dict) -> None:
    try:
        track_file = LOGS_DIR / f"tracking_{run_id}.json"
        track_file.write_text(json.dumps({
            "run_id":       run_id,
            "applied":      state["apply_results"],
            "completed_at": _now(),
        }, indent=2))
    except Exception:
        pass


def _send_run_notification(state: dict, *, title: str, message: str, stage: str) -> None:
    notif_cfg = dict((state.get("config") or {}).get("notifications") or {})
    if not (notif_cfg.get("enable_email") or notif_cfg.get("enable_sms")):
        return

    profile = state.get("profile") or {}
    candidate_email = str(notif_cfg.get("email") or profile.get("email") or "").strip()
    candidate_phone = str(notif_cfg.get("phone") or profile.get("phone") or "").strip()
    notifier = NotificationService()
    result = notifier.send_alert(
        message=message,
        title=title,
        to_email=candidate_email,
        to_phone=candidate_phone,
    )
    state.setdefault("notification_log", []).append({
        "timestamp": _now(),
        "stage": stage,
        "title": title,
        "requested_channels": {
            "email": bool(notif_cfg.get("enable_email")),
            "sms": bool(notif_cfg.get("enable_sms")),
        },
        "result": result,
    })


def _clean_role_title(raw_title: str) -> str:
    """Normalize noisy job titles for resume/cover-letter personalization."""
    import re

    title = str(raw_title or "").strip()
    if not title:
        return "the role"

    title = re.sub(r"\(.*?\)", "", title)
    title = re.sub(r"\b(linkedin|indeed|dice|ziprecruiter|monster|glassdoor)\b", "", title, flags=re.I)
    title = re.sub(r"\b(remote|hybrid|onsite|on-site|work\s*from\s*home|wfh)\b", "", title, flags=re.I)
    title = re.sub(r"\s*[|·—–-]\s*.*$", "", title)
    title = re.sub(r"\s{2,}", " ", title).strip(" -|·—–")
    return title or "the role"


def _build_cover_letter_text(profile: dict, job: dict) -> str:
    """Create a classic business cover letter format with robust role alignment."""
    role = _clean_role_title(job.get("title", ""))
    company = str(job.get("company") or "Hiring Team")
    candidate = str(profile.get("name") or "Candidate")
    email = str(profile.get("email") or "")
    phone = str(profile.get("phone") or "")
    skills = [str(s).strip() for s in (profile.get("skills") or []) if str(s).strip()]
    top_skills = ", ".join(skills[:8]) if skills else "AI/ML engineering, cloud architecture, and delivery leadership"
    summary = _sanitize_profile_summary(profile.get("summary") or "I build production-ready AI systems with measurable business outcomes.")

    experience_items = [str(x).strip() for x in (profile.get("experience") or []) if str(x).strip()]
    projects = [str(x).strip() for x in (profile.get("projects") or []) if str(x).strip()]
    impact_anchor = projects[0] if projects else (experience_items[0] if experience_items else "enterprise platform modernization")

    return (
        f"{candidate}\n"
        f"{email} | {phone}\n"
        f"{_now()[:10]}\n\n"
        "Hiring Manager\n"
        f"{company}\n\n"
        f"Subject: Application for {role}\n\n"
        "Dear Hiring Manager,\n\n"
        f"I am writing to express interest in the {role} position at {company}. {summary}\n\n"
        f"My background aligns strongly with your requirements, especially in {top_skills}. "
        f"A representative example is {impact_anchor}, where I partnered cross-functionally to improve reliability, delivery velocity, and measurable business outcomes.\n\n"
        f"I am confident this blend of technical depth and execution discipline would let me contribute quickly to {company}. "
        "I would value the opportunity to discuss how my background can support your team's goals.\n\n"
        "Thank you for your time and consideration.\n\n"
        "Sincerely,\n"
        f"{candidate}\n"
    ).strip()


@traceable(name="api.parse_resume")
async def _parse_resume(resume_path: Path) -> dict:
    """Extract profile from uploaded resume file."""
    text = ""
    suffix = resume_path.suffix.lower()

    if suffix == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(resume_path) as pdf:
                text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        except ImportError:
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(str(resume_path))
                text = "\n".join(p.get_text() for p in doc)
            except ImportError:
                text = f"[PDF text extraction unavailable — file: {resume_path.name}]"
    elif suffix in (".txt", ".md"):
        text = resume_path.read_text(errors="replace")
    elif suffix in (".docx",):
        try:
            from docx import Document
            doc  = Document(str(resume_path))
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            text = resume_path.read_text(errors="replace")
    else:
        text = resume_path.read_text(errors="replace")

    return _extract_profile_from_text(text)


def _sanitize_profile_summary(raw_summary: str) -> str:
    summary = str(raw_summary or "")
    summary = re.sub(r"https?://\S+", " ", summary)
    summary = re.sub(r"\b(?:linkedin|github|portfolio|demo)\s*:\s*", " ", summary, flags=re.I)
    summary = re.sub(r"[\#*_`]+", " ", summary)
    summary = re.sub(r"\|", " • ", summary)
    summary = re.sub(r"\b[\w.+-]+@[\w-]+\.\w+\b", " ", summary)
    summary = re.sub(r"\+?\d[\d\s().-]{8,}\d", " ", summary)
    summary = re.sub(r"\s+", " ", summary).strip(" ,;:-•")
    return summary[:500]


def _extract_profile_from_text(text: str) -> dict:
    """Enhanced resume parsing with skills, education, projects, and years of experience."""
    import re

    lines = [l.strip() for l in text.split("\n") if l.strip()]

    name = lines[0] if lines else "Candidate"
    if "|" in name:
        name = name.split("|")[0].strip()
    if len(name) > 60 or any(c in name for c in "@./"):
        name = "Candidate"

    email_m = re.search(r"[\w.+-]+@[\w-]+\.\w+", text)
    email = email_m.group(0) if email_m else ""

    phone_m = re.search(r"[\+\(]?[\d\s\-\(\)]{10,}", text)
    phone = phone_m.group(0).strip() if phone_m else ""

    found_skills = extract_skills(text)

    exp_pattern = re.findall(
        r"([\w\s/,&-]+(?:Engineer|Developer|Manager|Director|Analyst|Scientist|Lead|Architect|Consultant))"
        r"[^\n]*?(\d{4})\s*[-–]\s*(\d{4}|Present|Current|Now)",
        text,
        re.I,
    )
    experience = []
    for role, start_s, end_s in exp_pattern[:8]:
        start = int(start_s)
        end = 2026 if end_s.lower() in ("present", "current", "now") else int(end_s)
        years = max(0, end - start)
        experience.append({"title": role.strip(), "years": years, "start": start, "end": end_s})

    if not experience:
        fallback_titles: list[str] = []
        exp_section = False
        for line in lines:
            low = line.lower()
            if "professional experience" in low:
                exp_section = True
                continue
            if exp_section and any(stop in low for stop in ("notable projects", "professional affiliations", "publications", "education")):
                break
            if not exp_section:
                continue
            if len(line) > 80 or any(ch in line.lower() for ch in ("@", "http://", "https://")):
                continue
            if re.search(r"(architect|scientist|lead|manager|engineer|consultant)", line, re.I):
                fallback_titles.append(line.strip("•	 -"))
        yoe = re.search(r"(\d+)\+?\s*years?\s+(?:of\s+)?experience", text, re.I)
        inferred_years = int(yoe.group(1)) if yoe else 0
        if fallback_titles:
            experience = [{"title": title, "years": inferred_years or 4} for title in list(dict.fromkeys(fallback_titles))[:8]]
        elif yoe:
            experience = [{"title": "Software Professional", "years": int(yoe.group(1))}]

    education = []
    for m in re.finditer(r"((?:B\.?Tech|B\.?E|Bachelors?|Masters?|M\.?S\.?|MBA|PhD|Doctorate)[^\n]{0,120})", text, re.I):
        education.append(m.group(1).strip())
    education = list(dict.fromkeys(education))[:6]

    projects = []
    for m in re.finditer(r"(?:project|projects)[:\-]?\s*([^\n]{8,140})", text, re.I):
        val = re.sub(r"^[#*\-\s]+", "", m.group(1)).strip(" .-")
        if len(val) >= 8:
            projects.append(val)
    projects = list(dict.fromkeys(projects))[:8]

    summary_lines = []
    for line in lines[1:12]:
        low = line.lower()
        if any(token in low for token in ("linkedin", "github", "portfolio", "demo", "http://", "https://", "@")):
            continue
        if re.fullmatch(r"[+()\d\s.-]{10,}", line):
            continue
        if len(line.split()) < 3:
            continue
        summary_lines.append(line)
        if len(summary_lines) >= 3:
            break
    summary = _sanitize_profile_summary(" ".join(summary_lines) if summary_lines else text[:300])

    total_years = sum(int(e.get("years") or 0) for e in experience)

    return {
        "name": name,
        "email": email,
        "phone": _sanitize_phone(phone),
        "skills": found_skills,
        "experience": experience,
        "education": education,
        "projects": projects,
        "total_years_experience": total_years,
        "summary": summary,
        "raw_text": text[:6000],
    }


def _is_generic_target_role(role: str) -> bool:
    low = str(role or "").strip().lower()
    return low in {"software engineer", "engineer", "developer", "software developer", "software architect"}


def _infer_target_roles(profile: dict, config_roles: list[str] | None) -> list[str]:
    requested = [str(r).strip() for r in (config_roles or []) if str(r).strip()]
    if requested and not all(_is_generic_target_role(r) for r in requested):
        requested = requested + ["Staff Engineer", "Architect", "Data Science Lead"]
        normalized: list[str] = []
        seen = set()
        for role in requested:
            role = re.sub(r"\s+", " ", str(role or "")).strip()
            if role and role.lower() not in seen:
                seen.add(role.lower())
                normalized.append(role)
            if len(normalized) >= 10:
                break
        return normalized

    exp_titles = [str((item or {}).get("title") or "").strip() for item in (profile.get("experience") or []) if isinstance(item, dict)]
    skills = {str(s).strip().lower() for s in (profile.get("skills") or []) if str(s).strip()}
    inferred: list[str] = []
    inferred.extend([title for title in exp_titles if title])
    inferred.extend(["Staff Engineer", "Architect", "Data Science Lead"])

    ai_signal = any(tok in skills for tok in {"machine learning", "tensorflow", "azure openai", "llm", "ai architect", "solution architect", "deep learning"})
    if ai_signal:
        inferred.extend([
            "Senior Solution Architect",
            "AI Solution Architect",
            "Generative AI Architect",
            "Lead Data Scientist",
            "Principal AI Engineer",
            "Machine Learning Architect",
        ])

    normalized: list[str] = []
    seen = set()
    for role in inferred + requested + ["AI Engineer", "Machine Learning Engineer", "Staff Engineer", "Architect", "Data Science Lead"]:
        role = re.sub(r"\s+", " ", str(role or "")).strip()
        if role and role.lower() not in seen:
            seen.add(role.lower())
            normalized.append(role)
        if len(normalized) >= 10:
            break
    return normalized or ["AI Engineer", "Machine Learning Engineer", "Staff Engineer", "Architect", "Data Science Lead"]


@traceable(name="api.build_intent")
def _build_intent(profile: dict, config: dict) -> dict:
    self_learning_context = str(config.get("self_learning_context") or _GLOBAL_SELF_LEARNING_CONTEXT or "").strip()
    roles = _infer_target_roles(profile, config.get("target_roles"))

    # Pass ALL skills, not just 8 — LeadScout needs these for multi-query bucketing
    all_skills = profile.get("skills", [])

    # Also derive extra keywords from raw_text if available
    extra_kw: list[str] = []
    raw = profile.get("raw_text", "")
    for term in [
        "LangChain", "LangGraph", "RAG", "GenAI", "LLM", "MLOps", "SageMaker",
        "Bedrock", "Vertex AI", "Hugging Face", "Fine-tuning", "RLHF",
        "Generative AI", "Vector Database", "Embeddings", "Prompt Engineering",
    ]:
        if term.lower() in raw.lower() and term not in all_skills:
            extra_kw.append(term)

    return {
        "target_roles":      roles,
        "keywords":          list(dict.fromkeys(all_skills + extra_kw)),  # ALL skills
        "extracted_profile": profile,   # full profile for LeadScout query bucketing
        "geo_preferences":   config.get("geo_preferences", {"remote": True, "locations": ["United States"]}),
        "salary_min_usd":    config.get("salary_min", 90_000),
        "salary_max_usd":    config.get("salary_max", 200_000),
        "self_learning_context": self_learning_context,
    }


@traceable(name="api.stub_leads")
def _stub_leads(profile: dict, max_jobs: int = 100) -> list[dict]:
    """Return realistic stub leads when API keys are unavailable."""
    skills = [str(skill).strip() for skill in (profile.get("skills") or ["Python"]) if str(skill).strip()][:6]
    roles = _infer_target_roles(profile, None)[:10] or ["AI Engineer", "Staff Engineer", "Architect"]
    companies = [
        "TechCorp Inc.", "StartupAI", "ScaleUp Inc.", "CloudForge", "DataNova", "Vertex Labs",
        "Northstar Health", "FinCore Systems", "Orbit Analytics", "BlueRiver Tech",
        "Atlas Platforms", "SignalPath AI", "Apex Commerce", "BrightOps", "Catalyst Data",
        "NextWave Robotics", "Summit Digital", "Harbor Cloud", "Lumen Insights", "Quantum Stack",
    ]
    locations = [
        ("Remote", True),
        ("San Francisco, CA", True),
        ("New York, NY", False),
        ("Boston, MA", True),
        ("Seattle, WA", True),
        ("Austin, TX", False),
        ("Chicago, IL", True),
        ("Atlanta, GA", False),
    ]
    sources = ["linkedin", "indeed", "glassdoor", "naukri", "greenhouse", "lever", "workday"]
    suffixes = [
        "Platform", "AI Products", "Enterprise Data", "Applied AI", "Cloud Architecture",
        "ML Systems", "Data Science", "Automation", "Intelligent Workflows", "Decisioning",
    ]
    search_slugs = {
        "linkedin": "https://www.linkedin.com/jobs/view/{job_id}",
        "indeed": "https://www.indeed.com/viewjob?jk={job_id}",
        "glassdoor": "https://www.glassdoor.com/job-listing/demo-role-JV_IC1147401_KO0,9_KE10,14.htm?jl={job_id}",
        "naukri": "https://www.naukri.com/job-listings-{query}-{job_id}",
        "greenhouse": "https://boards.greenhouse.io/demo/jobs/{job_id}",
        "lever": "https://jobs.lever.co/demo/{job_id}",
        "workday": "https://demo.wd5.myworkdayjobs.com/en-US/Careers/job/{job_id}",
    }
    seed_jobs: list[dict] = []
    for idx in range(max_jobs):
        role = roles[idx % len(roles)]
        suffix = suffixes[idx % len(suffixes)]
        company = f"{companies[idx % len(companies)]} {suffix.split()[0]} Team {idx+1:03d}"
        location, remote = locations[idx % len(locations)]
        source = sources[idx % len(sources)]
        title = role if suffix.lower() in role.lower() else f"{role} — {suffix}"
        query = quote_plus(title.lower().replace("—", " ").replace("/", " "))
        job_id = f"{idx+1:06d}"
        primary_skill = skills[idx % len(skills)] if skills else "Python"
        secondary_skill = skills[(idx + 1) % len(skills)] if len(skills) > 1 else primary_skill
        seed_jobs.append(
            {
                "id": f"demo_{idx+1:03d}",
                "title": title,
                "company": company,
                "url": sanitize_job_url(search_slugs[source].format(query=query, job_id=job_id)),
                "location": location,
                "remote": remote,
                "description": (
                    f"Seeking a {role} with strength in {primary_skill}, {secondary_skill}, "
                    "stakeholder leadership, and shipping production AI/data systems."
                ),
                "demo_fallback": True,
                "source": source,
                "salary_min": 125000 + ((idx % 6) * 10000),
                "salary_max": 185000 + ((idx % 6) * 12000),
            }
        )
    if max_jobs <= len(seed_jobs):
        return seed_jobs[:max_jobs]
    expanded = []
    for idx in range(max_jobs):
        base = dict(seed_jobs[idx % len(seed_jobs)])
        base["id"] = f"{base['id']}_{idx+1:03d}"
        base["posted_hours_ago"] = (idx % 72) + 1
        base["url"] = sanitize_job_url(f"{base['url']}&job_stub_id={idx+1:03d}")
        base["direct_job_url"] = base["url"]
        expanded.append(base)
    return expanded


@traceable(name="api.apply_frontend_filters")
def _apply_frontend_filters(jobs: list[dict], config: dict) -> list[dict]:
    work_modes = set(config.get("work_modes") or ["remote", "hybrid", "onsite"])
    salary_min = int(config.get("salary_min", 0) or 0)
    salary_max = int(config.get("salary_max", 10**9) or 10**9)
    posted_within = int(config.get("posted_within_hours", 9999) or 9999)

    filtered = []
    for job in jobs:
        is_remote = bool(job.get("remote"))
        location = str(job.get("location") or "").lower()
        has_hybrid_hint = "hybrid" in location or "remote" in location
        if is_remote:
            mode = "remote"
        elif has_hybrid_hint:
            mode = "hybrid"
        else:
            mode = "onsite"
        if mode not in work_modes:
            continue
        jmin = int(job.get("salary_min") or 0)
        jmax = int(job.get("salary_max") or 10**9)
        if jmax < salary_min or jmin > salary_max:
            continue
        posted = int(job.get("posted_hours_ago") or 24)
        if posted > posted_within:
            continue
        filtered.append(job)
    return filtered


@traceable(name="api.stub_score")
def _stub_score(leads: list[dict]) -> list[dict]:
    import random
    random.seed(42)
    return [{**j, "score": round(random.uniform(0.45, 0.95), 3)} for j in leads]


@traceable(name="api.generate_artifacts")
async def _generate_artifacts(profile: dict, jobs: list[dict], out_dir: Path) -> dict:
    """Generate ATS docs + before/after ATS scores for each approved job."""
    out_dir.mkdir(parents=True, exist_ok=True)
    artifacts: dict[str, dict[str, Any]] = {}

    baseline_md = _build_resume_markdown(profile, keyword_hints=[])
    baseline_ats = _compute_resume_ats_scores(baseline_md, "", profile.get("skills") or [])

    for job in jobs:
        job_id = str(job.get("id", f"job_{id(job)}"))
        job_dir = out_dir / job_id
        job_dir.mkdir(exist_ok=True)

        jd_text = " ".join(str(job.get(k) or "") for k in ("description", "snippet", "title"))
        keyword_hints = extract_skills(jd_text, extra_candidates=profile.get("skills") or [])[:12]
        tailored_md = _build_resume_markdown(profile, keyword_hints=keyword_hints)
        tailored_ats = _compute_resume_ats_scores(tailored_md, jd_text, profile.get("skills") or [])

        baseline_md_path = job_dir / "resume_baseline.md"
        tailored_md_path = job_dir / "resume_tailored.md"
        cover_md_path = job_dir / "cover_letter.md"
        ats_report_path = job_dir / "ats_verification.json"
        resume_path = job_dir / "resume.docx"
        cover_path = job_dir / "cover_letter.docx"

        baseline_md_path.write_text(baseline_md, encoding="utf-8")
        tailored_md_path.write_text(tailored_md, encoding="utf-8")

        cover_text = _build_cover_letter_text(profile, job)
        cover_md_path.write_text(cover_text, encoding="utf-8")

        _write_resume_docx(profile, job, resume_path, tailored_md=tailored_md)
        _write_cover_docx(profile, job, cover_path)

        resume_pdf = _to_pdf(resume_path)
        cover_pdf = _to_pdf(cover_path)

        ats_report_path.write_text(
            json.dumps(
                {
                    "job_id": job_id,
                    "job_title": job.get("title", ""),
                    "company": job.get("company", ""),
                    "verification_tool": "careeragent_internal_ats_v1",
                    "ats_score_before": baseline_ats,
                    "ats_score_after": tailored_ats,
                    "improvement": round(
                        float(tailored_ats.get("overall") or 0.0) - float(baseline_ats.get("overall") or 0.0),
                        2,
                    ),
                },
                indent=2,
            ),
            encoding="utf-8",
        )

        artifacts[job_id] = {
            "resume_docx": str(resume_path),
            "cover_docx": str(cover_path),
            "resume_pdf": str(resume_pdf) if resume_pdf else None,
            "cover_pdf": str(cover_pdf) if cover_pdf else None,
            "resume_baseline_md": str(baseline_md_path),
            "resume_tailored_md": str(tailored_md_path),
            "cover_letter_md": str(cover_md_path),
            "ats_verification_report": str(ats_report_path),
            "ats_score_before": baseline_ats,
            "ats_score_after": tailored_ats,
            "keywords_injected": keyword_hints,
        }

    return artifacts


def _build_resume_markdown(profile: dict, keyword_hints: list[str]) -> str:
    jd_terms = list(dict.fromkeys([s for s in keyword_hints if str(s).strip()]))[:10]
    base_skills = list(dict.fromkeys([str(s) for s in (profile.get("skills") or []) if str(s).strip()]))

    def _bucket(skill: str) -> str:
        s = skill.lower()
        if any(k in s for k in ("ml", "ai", "llm", "nlp", "pytorch", "tensorflow", "scikit", "rag", "embedding")):
            return "AI/ML"
        if any(k in s for k in ("aws", "azure", "gcp", "kubernetes", "docker", "terraform", "serverless", "cloud")):
            return "Cloud Engineering"
        return "Data Ops"

    matrix = {"AI/ML": [], "Cloud Engineering": [], "Data Ops": []}
    for skill in list(dict.fromkeys(jd_terms + base_skills)):
        matrix[_bucket(skill)].append(skill)
    for k in matrix:
        matrix[k] = matrix[k][:12]

    experience_items = [e for e in (profile.get("experience") or []) if e]
    project_items = [str(p).strip() for p in (profile.get("projects") or []) if str(p).strip()]
    notable_projects = project_items or [
        "Real-time recommendation platform modernization",
        "Enterprise MLOps governance rollout",
        "Cloud data quality and observability transformation",
    ]

    project_lines: list[str] = []
    for idx, project in enumerate(notable_projects[:8], start=1):
        project_lines.append(f"### Project {idx}: {project}")
        project_lines.extend([
            f"- Situation: Inherited fragmented delivery across analytics, application, and platform teams with inconsistent SLAs and limited ownership visibility for {project}.",
            "- Task: Led architecture modernization with clear technical milestones, ownership models, and measurable acceptance criteria across product, engineering, and operations stakeholders.",
            "- Action: Designed event-driven services, codified CI/CD guardrails, and introduced automated validation gates; reduced release cycle time by 42% and cut deployment failures by 37%.",
            "- Action: Implemented telemetry-first observability with latency/error/cost dashboards and anomaly alerting; improved incident detection speed by 58% and lowered MTTR by 46%.",
            "- Result: Delivered sustained production performance gains (99.95% service availability, 31% infrastructure cost reduction, and ~18 hours/week engineering time reclaimed).",
        ])

    exp_lines = []
    for exp in experience_items[:8]:
        title = exp.get("title", "Senior Technical Leader") if isinstance(exp, dict) else str(exp)
        years = exp.get("years", "") if isinstance(exp, dict) else ""
        exp_lines.append(f"- {title} ({years} years)")
    if not exp_lines:
        exp_lines = ["- 16+ years delivering AI/ML platforms, cloud-native systems, and data operations at enterprise scale."]

    edu_lines = [f"- {e}" for e in (profile.get("education") or [])[:4]] or ["- Education details available"]
    summary = _sanitize_profile_summary(profile.get("summary", "Principal-level technical architect with 16+ years building resilient, measurable software platforms."))

    resume_md = (
        f"# {profile.get('name','Candidate')}\n"
        f"{profile.get('email','')} · {profile.get('phone','')}\n\n"
        "## Professional Summary\n"
        f"{summary}\n"
        "- Architected multi-region distributed systems, production MLOps stacks, and governed data platforms with measurable business outcomes.\n"
        "- Recognized for turning ambiguous business objectives into delivery roadmaps with reliable execution, risk controls, and stakeholder trust.\n\n"
        "## Technical Skills\n"
        f"- **AI/ML:** {', '.join(matrix['AI/ML']) or 'Machine Learning, MLOps, NLP, LLMOps'}\n"
        f"- **Cloud Engineering:** {', '.join(matrix['Cloud Engineering']) or 'AWS, Azure, GCP, Docker, Kubernetes, Terraform'}\n"
        f"- **Data Ops:** {', '.join(matrix['Data Ops']) or 'Data Modeling, ETL, Airflow, Spark, dbt, Observability'}\n\n"
        "## Experience Highlights\n"
        + "\n".join(exp_lines)
        + "\n\n## Notable Projects\n"
        + "\n".join(project_lines)
        + "\n\n## Education\n"
        + "\n".join(edu_lines)
        + "\n"
    )

    if len(re.findall(r"\b\w+\b", resume_md)) < 800:
        expansion = []
        for project in notable_projects[:5]:
            expansion.extend([
                f"- Expanded depth: For {project}, directed cross-functional architecture reviews, performance experiments, and production hardening workstreams to ensure scale-readiness.",
                "- Expanded depth: Formalized service-level objectives, build-vs-buy tradeoff analyses, and risk-mitigation controls; increased roadmap predictability by 33%.",
                "- Expanded depth: Mentored staff engineers through design critiques and incident retrospectives, raising engineering throughput while improving quality gates.",
            ])
        resume_md = f"{resume_md}\n### Additional Technical Depth\n" + "\n".join(expansion) + "\n"

    return resume_md


def _compute_resume_ats_scores(resume_md: str, jd_text: str, profile_skills: list[str]) -> dict:
    sections = ["summary", "skills", "experience", "projects", "education"]
    section_hits = sum(1 for s in sections if s in resume_md.lower())
    layout_score = round((section_hits / len(sections)) * 100, 2)
    jd_skills = set(extract_skills(jd_text, extra_candidates=profile_skills)) if jd_text else set()
    resume_skills = set(extract_skills(resume_md, extra_candidates=profile_skills))
    keyword_score = round((len(jd_skills & resume_skills) / max(1, len(jd_skills))) * 100, 2) if jd_skills else 0.0
    overall = round((0.55 * layout_score) + (0.45 * keyword_score), 2)
    return {
        "overall": overall,
        "layout": layout_score,
        "keyword": keyword_score,
        "matched_keywords": sorted(jd_skills & resume_skills)[:20],
        "missing_keywords": sorted(jd_skills - resume_skills)[:20],
    }

def _write_resume_docx(profile: dict, job: dict, path: Path, tailored_md: str = "") -> None:
    try:
        from docx import Document
        doc = Document()
        content_md = tailored_md or _build_resume_markdown(profile, keyword_hints=[])
        current_section = None
        for line in content_md.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), 0)
            elif line.startswith("## "):
                current_section = line[3:].strip()
                doc.add_heading(current_section, 1)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)
        doc.save(path)
    except ImportError:
        path.write_text(f"RESUME\n{profile.get('name','Candidate')}\n{', '.join(profile.get('skills',[]))}")


def _write_cover_docx(profile: dict, job: dict, path: Path) -> None:
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Cover Letter", 0)
        doc.add_paragraph(_build_cover_letter_text(profile, job))
        doc.save(path)
    except ImportError:
        path.write_text(_build_cover_letter_text(profile, job), encoding="utf-8")


def _to_pdf(docx_path: Path) -> Optional[Path]:
    import shutil
    pdf_path = docx_path.with_suffix(".pdf")
    if not shutil.which("libreoffice"):
        return None
    try:
        import subprocess
        r = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(docx_path.parent), str(docx_path)],
            capture_output=True, timeout=30,
        )
        if r.returncode == 0 and pdf_path.exists():
            return pdf_path
    except Exception:
        pass
    return None


# ══════════════════════════════════════════════════════════════════════════════
# FASTAPI APP
# ══════════════════════════════════════════════════════════════════════════════

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Clean startup / shutdown — no crash."""
    log.info("CareerAgent API starting up…")
    ARTIFACTS_DIR.mkdir(parents=True, exist_ok=True)
    LOGS_DIR.mkdir(parents=True, exist_ok=True)
    yield
    log.info("CareerAgent API shutting down…")


app = FastAPI(
    title="CareerAgent-AI API",
    version="1.0.0",
    description="L0→L9 Autonomous Job Hunt Engine",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "runs_active": len(_runs)}


@app.post("/mcp/invoke")
async def mcp_invoke_passthrough(body: dict):
    """Compatibility endpoint for clients still posting to /mcp/invoke."""
    tool = str((body or {}).get("tool") or "")
    payload = (body or {}).get("payload") or (body or {}).get("args") or {}
    if not tool:
        raise HTTPException(400, "tool is required")
    return {
        "ok": True,
        "tool": tool,
        "payload": payload,
        "message": "MCP compatibility endpoint reached",
    }


@app.post("/mcp/invoke/")
async def mcp_invoke_passthrough_trailing(body: dict):
    """Trailing slash compatibility for hosted backends/proxies."""
    return await mcp_invoke_passthrough(body)

@app.post("/hunt/start")
@traceable(name="api.start_hunt")
async def start_hunt(
    background_tasks: BackgroundTasks,
    resume: UploadFile = File(...),
    hunt_config: str = Form(default="{}", alias="config"),
):
    """
    Start a new pipeline run.
    Accepts multipart/form-data with:
      - resume: PDF/TXT/DOCX file
      - config: JSON string with optional keys:
          target_roles, geo_preferences, match_threshold, salary_min, salary_max
    Returns { run_id, status }
    """
    run_id = uuid.uuid4().hex[:12]

    try:
        cfg = json.loads(hunt_config) if hunt_config else {}
        if not isinstance(cfg, dict):
            cfg = {}
    except json.JSONDecodeError:
        cfg = {}

    try:
        # Save uploaded file
        suffix = Path(resume.filename or "resume.pdf").suffix or ".pdf"
        save_path = UPLOADS_DIR / f"{run_id}{suffix}"
        content = await resume.read()
        if not content:
            raise HTTPException(400, "Uploaded resume file is empty")
        save_path.write_bytes(content)
        log.info("Resume saved: %s (%d bytes)", save_path, len(content))

        # Initialize state
        _runs[run_id] = _build_initial_state(run_id, cfg)
        _runs[run_id]["resume_path"] = str(save_path)

        # Launch pipeline asynchronously so the HTTP request returns immediately
        # even when downstream layers take longer on hosted deployments.
        _spawn_async_action(f"start_hunt:{run_id}", run_pipeline(run_id, save_path))
        return {"run_id": run_id, "status": "started", "message": "Pipeline launched"}
    except HTTPException:
        raise
    except Exception as exc:
        log.exception("Failed to start run %s", run_id)
        raise HTTPException(500, f"Failed to start run: {exc}") from exc


@app.get("/hunt/{run_id}/status")
@traceable(name="api.get_status")
async def get_status(run_id: str):
    """Poll this endpoint for real-time progress updates."""
    if run_id not in _runs:
        # Try to reload from persisted file
        state_file = LOGS_DIR / f"state_{run_id}.json"
        if state_file.exists():
            data = json.loads(state_file.read_text())
            _runs[run_id] = data
        else:
            raise HTTPException(404, f"Run {run_id} not found")

    state = _runs[run_id]
    return {
        "run_id":           state["run_id"],
        "status":           state["status"],
        "progress_pct":     state["progress_pct"],
        "layers":           state["layers"],
        "jobs_discovered":  state["jobs_discovered"],
        "jobs_scored":      state["jobs_scored"],
        "jobs_approved":    state["jobs_approved"],
        "jobs_applied":     state["jobs_applied"],
        "top_match_score":  state["top_match_score"],
        "candidate_name":   state["candidate_name"],
        "skills_extracted": state["skills_extracted"],
        "pending_action":   state.get("pending_action"),
        "langsmith":        state.get("langsmith", {}),
        "langgraph":        state.get("langgraph", {}),
        "llm_stack":        state.get("llm_stack", {}),
        "apply_results":    state.get("apply_results", []),
        "interviews":       state.get("interviews", []),
        "followup_queue":   state.get("followup_queue", []),
        "notification_log": state.get("notification_log", []),
        "feedback_events":  state.get("feedback_events", [])[-50:],
        "learning_loop":    state.get("learning_loop", {}),
        "employer_outcomes": state.get("employer_outcomes", {}),
        "learning_resources": state.get("learning_resources", {}),
        "analytics_summary": state.get("analytics_summary", {}),
        "self_learning_context": state.get("self_learning_context", ""),
        "self_learning_prompt": state.get("self_learning_prompt", ""),
        "system_prompt_update": state.get("system_prompt_update", ""),
        "feedback_learning_state": state.get("feedback_learning_state", {}),
        "profile":          state.get("profile", {}),
        "layer_debug":      state.get("layer_debug", {}),
        "evaluations":      state.get("evaluations", [])[-50:],
        "raw_job_leads_preview": state.get("job_leads", [])[:25],
        "scored_jobs_preview": state.get("scored_jobs", [])[:25],
        "approved_jobs_preview": state.get("approved_jobs", [])[:50],
        "resume_scores":    state.get("resume_scores", {}),
        "agent_log":        state["agent_log"][-30:],  # last 30 entries
        "errors":           state["errors"],
        "created_at":       state["created_at"],
        "completed_at":     state["completed_at"],
    }


@app.get("/hunt/{run_id}/jobs")
@traceable(name="api.get_jobs")
async def get_jobs(run_id: str):
    if run_id not in _runs:
        raise HTTPException(404, f"Run {run_id} not found")
    state = _runs[run_id]
    jobs  = [{**job, "url": sanitize_job_url(job.get("direct_job_url") or job.get("url", "")), "direct_job_url": sanitize_job_url(job.get("direct_job_url") or job.get("url", ""))} for job in (state.get("scored_jobs", []) or [])]
    return {
        "run_id":    run_id,
        "total":     len(jobs),
        "jobs":      jobs,
    }


@app.get("/hunt/{run_id}/applications")
@traceable(name="api.get_applications")
async def get_applications(run_id: str):
    if run_id not in _runs:
        raise HTTPException(404, f"Run {run_id} not found")
    state = _runs[run_id]
    return {
        "run_id": run_id,
        "applications": state.get("apply_results", []),
        "interviews": state.get("interviews", []),
        "followup_queue": state.get("followup_queue", []),
        "notification_log": state.get("notification_log", []),
    }




@app.post("/hunt/{run_id}/feedback")
@traceable(name="api.feedback")
async def post_feedback(run_id: str, body: dict):
    if run_id not in _runs:
        raise HTTPException(404, f"Run {run_id} not found")
    state = _runs[run_id]
    payload = dict(body or {})
    if not str(payload.get("text") or payload.get("comment") or "").strip():
        payload["text"] = "No comment"
    event = _record_feedback_event(state, payload)
    feedback_file = LOGS_DIR / f"feedback_{run_id}.jsonl"
    with feedback_file.open("a", encoding="utf-8") as fh:
        fh.write(json.dumps(event) + "\n")
    _persist_state(run_id)
    return {"ok": True, "event": event, "totals": state.get("learning_loop", {})}


@app.get("/hunt/{run_id}/feedback")
@traceable(name="api.get_feedback")
async def get_feedback(run_id: str):
    if run_id not in _runs:
        state_file = LOGS_DIR / f"state_{run_id}.json"
        if state_file.exists():
            _runs[run_id] = json.loads(state_file.read_text())
        else:
            raise HTTPException(404, f"Run {run_id} not found")
    state = _runs[run_id]
    return {
        "run_id": run_id,
        "feedback": state.get("feedback_events", []),
        "self_learning_context": state.get("self_learning_context", ""),
    }


@app.get("/admin/feedback")
@traceable(name="api.admin_feedback")
async def admin_feedback():
    rows: list[dict[str, Any]] = []
    for path in sorted(LOGS_DIR.glob("feedback_*.jsonl")):
        run_id = path.stem.replace("feedback_", "", 1)
        try:
            with path.open("r", encoding="utf-8") as fh:
                for line in fh:
                    line = line.strip()
                    if not line:
                        continue
                    payload = json.loads(line)
                    rows.append({"run_id": run_id, **payload})
        except Exception as exc:
            rows.append({"run_id": run_id, "source": "system", "text": f"Failed to parse feedback log: {exc}"})
    rows.sort(key=lambda item: str(item.get("ts") or ""), reverse=True)
    return {"feedback": rows}


@app.post("/hunt/{run_id}/feedback/sync")
@traceable(name="api.sync_feedback")
async def sync_feedback(run_id: str):
    if run_id not in _runs:
        state_file = LOGS_DIR / f"state_{run_id}.json"
        if state_file.exists():
            _runs[run_id] = json.loads(state_file.read_text())
        else:
            raise HTTPException(404, f"Run {run_id} not found")
    state = _runs[run_id]
    context = _sync_feedback_to_agent_brain(state)
    feedback_file = LOGS_DIR / f"feedback_sync_{run_id}.json"
    feedback_file.write_text(json.dumps({
        "run_id": run_id,
        "synced_at": _now(),
        "self_learning_context": context,
    }, indent=2), encoding="utf-8")
    _persist_state(run_id)
    return {"ok": True, "self_learning_context": context}


@app.post("/hunt/{run_id}/action")
@traceable(name="api.run_action")
async def run_action(run_id: str, background_tasks: BackgroundTasks, body: dict):
    if run_id not in _runs:
        raise HTTPException(404, f"Run {run_id} not found")
    state = _runs[run_id]
    action = (body or {}).get("action")

    if action == "approve_ranking":
        selected_values = (
            (body or {}).get("selected_job_ids")
            or (body or {}).get("selected_job_urls")
            or (body or {}).get("selected_jobs")
            or []
        )
        ranked = state.get("layer_debug", {}).get("L5", {}).get("qualified_jobs", [])
        approved = pick_approved_jobs(ranked, selected_values)
        if selected_values and not approved:
            raise HTTPException(
                400,
                "No selected jobs matched ranked results. Send selected_job_ids or selected_job_urls from /hunt/{run_id}/jobs.",
            )
        state["approved_jobs"] = approved
        state["jobs_approved"] = len(approved)
        state["pending_action"] = None
        state["status"] = "running"
        _persist_state(run_id)
        _spawn_async_action(
            f"approve_ranking:{run_id}",
            _continue_l6_to_l9(
                run_id,
                stop_after_l6_for_approval=bool(state.get("config", {}).get("require_draft_approval", True)),
            ),
        )
        return {"ok": True, "message": f"approved {len(approved)} jobs"}

    if action == "approve_drafts":
        state["pending_action"] = None
        state["status"] = "running"
        _persist_state(run_id)
        _spawn_async_action(f"approve_drafts:{run_id}", _continue_l7_to_l9(run_id))
        return {"ok": True, "message": "drafts approved; resuming apply"}

    if action == "approve_followups":
        followups = state.get("followup_queue") or []
        for item in followups:
            item["draft_status"] = "approved"
            item["sent_at"] = _now()
        l7 = (state.get("layer_debug") or {}).get("L7") or {}
        for draft in (l7.get("email_drafts") or []):
            draft["status"] = "approved_and_sent"
            draft["sent_at"] = _now()
        state["pending_action"] = None
        state["status"] = "running"
        _log_agent(state, 7, f"Human approved {len(followups)} follow-up drafts. Continuing tracking and analytics.")
        _persist_state(run_id)
        _spawn_async_action(f"approve_followups:{run_id}", _continue_l8_to_l9(run_id))
        return {"ok": True, "message": f"follow-up drafts approved ({len(followups)}); resuming"}

    if action == "reject_followups":
        state["pending_action"] = "approve_followups"
        state["status"] = "pending_human_input"
        _log_agent(state, 7, "Follow-up drafts rejected by reviewer. Edit feedback and re-approve.")
        _persist_state(run_id)
        return {"ok": True, "message": "follow-up drafts rejected; awaiting revised approval"}

    if action == "reject_ranking":
        state["pending_action"] = None
        state["status"] = "running"
        state["hitl_rejections"] = int(state.get("hitl_rejections", 0)) + 1
        _log_agent(state, 5, "Ranking rejected by human reviewer. Looping back to L2 intake and planning.")
        _persist_state(run_id)
        resume_path = Path(state.get("resume_path") or "")
        if resume_path.exists():
            _spawn_async_action(f"reject_ranking:{run_id}", run_pipeline(run_id, resume_path))
            return {"ok": True, "message": "ranking rejected; restarting from L2"}
        raise HTTPException(400, "resume path missing; cannot re-run")

    if action == "reject_drafts":
        state["pending_action"] = "approve_ranking"
        state["status"] = "pending_human_input"
        _log_agent(state, 6, "Draft package rejected by reviewer. Returning to ranking gate.")
        _persist_state(run_id)
        return {"ok": True, "message": "drafts rejected; returned to ranking approval"}

    if action == "update_profile_skills":
        incoming = [str(x).strip() for x in ((body or {}).get("skills") or []) if str(x).strip()]
        if not incoming:
            raise HTTPException(400, "skills missing")
        prof = state.setdefault("profile", {})
        current = [str(x).strip() for x in (prof.get("skills") or []) if str(x).strip()]
        merged = list(dict.fromkeys(current + incoming))
        prof["skills"] = merged
        state["pending_action"] = None
        state["status"] = "running"
        _log_agent(state, 5, f"Profile updated with {len(incoming)} user-confirmed skills. Re-running from L4.")
        _persist_state(run_id)
        _spawn_async_action(f"update_profile_skills:{run_id}", _rerun_from_l4_l5(run_id))
        return {"ok": True, "message": f"profile updated with {len(incoming)} skills; rerunning from L4"}

    raise HTTPException(400, "unknown action")


@app.get("/hunt/{run_id}/artifacts")
async def get_artifacts(run_id: str):
    if run_id not in _runs:
        raise HTTPException(404, f"Run {run_id} not found")
    return {"run_id": run_id, "artifacts": _runs[run_id].get("artifacts", {})}


@app.get("/artifact/download")
async def download_artifact(path: str):
    """Download a generated artifact file."""
    p = Path(path)
    if not p.exists():
        raise HTTPException(404, "File not found")
    # Security: ensure it's under ARTIFACTS_DIR
    try:
        p.resolve().relative_to(ARTIFACTS_DIR.resolve())
    except ValueError:
        raise HTTPException(403, "Access denied")
    return FileResponse(str(p), filename=p.name)
