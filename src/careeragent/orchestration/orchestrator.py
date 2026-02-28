"""
Master Orchestrator — CareerAgent Pipeline L0 → L9
====================================================
Handles the full run lifecycle with Safe Method Resolution,
stateful progress tracking, artifact generation, and Playwright automation.

Layer Map:
  L0  — Profile Ingestion (up to 4000-token chunking)
  L1  — Profile Parsing & Validation
  L2  — Intent Planning (role targets, salary bands, geo prefs)
  L3  — LeadScout   (job discovery via search APIs + Playwright scrape)
  L4  — GeoFence    (location / remote filter)
  L5  — Extraction  (JD → structured requirements + match scoring)
  L6  — Artifact Generation (ATS resume + cover letter → .pdf / .docx)
  L7  — ApplyExecutor (Playwright form-fill + file upload)
  L8  — Confirmation & Receipt capture
  L9  — State flush + UI Progress Bar sync
"""

from __future__ import annotations

import asyncio
import inspect
import json
import logging
import os
import traceback
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Optional

# ── Logging ──────────────────────────────────────────────────────────────────
LOG_DIR = Path("logs")
LOG_DIR.mkdir(exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
    handlers=[
        logging.FileHandler(LOG_DIR / "careeragent.log"),
        logging.StreamHandler(),
    ],
)
log = logging.getLogger("orchestrator")

ARTIFACTS_DIR = Path("artifacts")
ARTIFACTS_DIR.mkdir(exist_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
# AGENT STATE  (single source of truth for UI Progress Bar)
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class LayerStatus:
    layer: int
    name: str
    status: str = "pending"          # pending | running | ok | error | skipped
    started_at: Optional[str] = None
    finished_at: Optional[str] = None
    error: Optional[str] = None
    meta: dict = field(default_factory=dict)


@dataclass
class AgentState:
    run_id: str = field(default_factory=lambda: datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S"))
    layers: list[LayerStatus] = field(default_factory=lambda: [
        LayerStatus(i, name) for i, name in enumerate([
            "Profile Ingestion",        # L0
            "Profile Parsing",          # L1
            "Intent Planning",          # L2
            "LeadScout",                # L3
            "GeoFence",                 # L4
            "JD Extraction & Scoring",  # L5
            "Artifact Generation",      # L6
            "ApplyExecutor",            # L7
            "Confirmation Capture",     # L8
            "State Flush",              # L9
        ])
    ])
    raw_profile_text: str = ""
    extracted_profile: dict = field(default_factory=dict)
    intent_plan: dict = field(default_factory=dict)
    job_leads: list[dict] = field(default_factory=list)
    filtered_leads: list[dict] = field(default_factory=list)
    scored_jobs: list[dict] = field(default_factory=list)
    artifacts: dict = field(default_factory=dict)   # {job_id: {"resume": path, "cover": path}}
    apply_results: list[dict] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    # ── helpers ──────────────────────────────────────────────────────────────
    def mark_running(self, layer: int) -> None:
        ls = self.layers[layer]
        ls.status = "running"
        ls.started_at = _now()
        self._persist()

    def mark_ok(self, layer: int, **meta) -> None:
        ls = self.layers[layer]
        ls.status = "ok"
        ls.finished_at = _now()
        ls.meta.update(meta)
        self._persist()

    def mark_error(self, layer: int, err: str) -> None:
        ls = self.layers[layer]
        ls.status = "error"
        ls.finished_at = _now()
        ls.error = err
        self.errors.append(f"L{layer}: {err}")
        self._persist()

    def progress_pct(self) -> float:
        done = sum(1 for ls in self.layers if ls.status in ("ok", "skipped", "error"))
        return round(done / len(self.layers) * 100, 1)

    def _persist(self) -> None:
        state_file = LOG_DIR / f"state_{self.run_id}.json"
        state_file.write_text(json.dumps(asdict(self), indent=2, default=str))
        log.debug("State persisted → %s  (%.0f%%)", state_file.name, self.progress_pct())


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


# ══════════════════════════════════════════════════════════════════════════════
# SAFE METHOD RESOLVER  (L3-L5 guard — no method-name crash)
# ══════════════════════════════════════════════════════════════════════════════

class SafeMethodResolver:
    """
    Resolves a callable on a service object by name OR by a ranked list of
    fallback aliases.  Returns a no-op coroutine if nothing matches, so the
    pipeline never crashes on a missing method.
    """

    ALIASES: dict[str, list[str]] = {
        # LeadScout
        "search_jobs":       ["search_jobs", "find_jobs", "scrape_jobs", "run", "execute"],
        # GeoFence
        "filter_by_geo":     ["filter_by_geo", "apply_geo_filter", "geo_filter", "filter"],
        # Extraction
        "extract_and_score": ["extract_and_score", "score_jobs", "extract_jd", "score", "run"],
        # ArtifactGenerator
        "generate_artifacts":["generate_artifacts", "create_artifacts", "build_docs", "run"],
        # ApplyExecutor
        "apply_to_job":      ["apply_to_job", "submit_application", "apply", "run"],
    }

    @classmethod
    def resolve(cls, service: Any, canonical: str) -> Callable:
        candidates = cls.ALIASES.get(canonical, [canonical])
        for name in candidates:
            method = getattr(service, name, None)
            if callable(method):
                log.debug("SafeMethodResolver: %s.%s → '%s'",
                          type(service).__name__, canonical, name)
                return method
        # Fallback — returns empty coroutine so pipeline continues
        log.warning(
            "SafeMethodResolver: No method found for '%s' on %s — using no-op",
            canonical, type(service).__name__,
        )
        async def _noop(*args, **kwargs):
            return []
        return _noop


# ══════════════════════════════════════════════════════════════════════════════
# PROFILE CHUNKER  (handles 4000-token profiles)
# ══════════════════════════════════════════════════════════════════════════════

MAX_TOKENS_PER_CHUNK = 3800   # safe margin below 4096
APPROX_CHARS_PER_TOKEN = 4


def chunk_profile(text: str) -> list[str]:
    """Split raw profile text into ≤4000-token chunks for LLM processing."""
    max_chars = MAX_TOKENS_PER_CHUNK * APPROX_CHARS_PER_TOKEN
    if len(text) <= max_chars:
        return [text]
    chunks, start = [], 0
    while start < len(text):
        end = start + max_chars
        # Try to break on a paragraph boundary
        boundary = text.rfind("\n\n", start, end)
        if boundary == -1 or boundary <= start:
            boundary = text.rfind("\n", start, end)
        if boundary == -1 or boundary <= start:
            boundary = end
        chunks.append(text[start:boundary].strip())
        start = boundary
    log.info("Profile chunked into %d piece(s) (total chars=%d)", len(chunks), len(text))
    return chunks


# ══════════════════════════════════════════════════════════════════════════════
# MASTER ORCHESTRATOR
# ══════════════════════════════════════════════════════════════════════════════

class Orchestrator:
    """
    Drives the L0 → L9 pipeline.

    Dependencies are injected so each service can be mocked or swapped:
        orchestrator = Orchestrator(
            profile_parser   = MyProfileParser(),
            intent_planner   = MyIntentPlanner(),
            lead_scout       = LeadScoutService(),
            geo_fence        = GeoFenceManager(),
            jd_extractor     = ExtractionManager(),
            artifact_gen     = ArtifactGenerator(),
            apply_executor   = ApplyExecutor(),
        )
    """

    def __init__(
        self,
        profile_parser=None,
        intent_planner=None,
        lead_scout=None,
        geo_fence=None,
        jd_extractor=None,
        artifact_gen=None,
        apply_executor=None,
        match_threshold: float = 0.45,   # lowered from 0.55 per debug checklist
    ):
        self.profile_parser = profile_parser
        self.intent_planner = intent_planner
        self.match_threshold = match_threshold

        # Auto-wire concrete services when callers don't inject dependencies.
        self.lead_scout = lead_scout
        if self.lead_scout is None:
            try:
                from careeragent.managers.leadscout_service import LeadScoutService

                self.lead_scout = LeadScoutService(enable_playwright_scrape=False)
            except Exception as exc:
                log.warning("LeadScout auto-wiring failed: %s", exc)

        self.geo_fence = geo_fence
        self.jd_extractor = jd_extractor
        if self.geo_fence is None or self.jd_extractor is None:
            try:
                from careeragent.managers.managers import GeoFenceManager, ExtractionManager

                self.geo_fence = self.geo_fence or GeoFenceManager()
                self.jd_extractor = self.jd_extractor or ExtractionManager()
            except Exception as exc:
                log.warning("GeoFence/Extraction auto-wiring failed: %s", exc)

        self.artifact_gen = artifact_gen
        self.apply_executor = apply_executor or ApplyExecutor()

    # ── Public entry point ───────────────────────────────────────────────────

    async def run(self, raw_profile: str, job_targets: Optional[list[str]] = None) -> AgentState:
        state = AgentState()
        state.raw_profile_text = raw_profile
        log.info("═══ Run %s started ═══", state.run_id)

        pipeline = [
            (0, self._l0_ingest),
            (1, self._l1_parse),
            (2, self._l2_plan),
            (3, self._l3_lead_scout),
            (4, self._l4_geo_fence),
            (5, self._l5_extract_score),
            (6, self._l6_generate_artifacts),
            (7, self._l7_apply),
            (8, self._l8_confirm),
            (9, self._l9_flush),
        ]

        for layer_num, handler in pipeline:
            state.mark_running(layer_num)
            try:
                await handler(state, job_targets=job_targets)
                log.info("✓ L%d %s  (%.0f%% complete)",
                         layer_num, state.layers[layer_num].name, state.progress_pct())
            except Exception as exc:
                tb = traceback.format_exc()
                log.error("✗ L%d FAILED:\n%s", layer_num, tb)
                state.mark_error(layer_num, str(exc))
                # Non-fatal layers: continue pipeline; fatal layers: abort
                if layer_num in (0, 1):
                    log.critical("Fatal layer failed — aborting run.")
                    break

        log.info("═══ Run %s finished  (%.0f%% layers OK) ═══",
                 state.run_id, state.progress_pct())
        return state

    # ── Layer Handlers ────────────────────────────────────────────────────────

    async def _l0_ingest(self, state: AgentState, **_):
        """L0 — Profile Ingestion with 4000-token chunking."""
        chunks = chunk_profile(state.raw_profile_text)
        state.mark_ok(0, chunks=len(chunks), total_chars=len(state.raw_profile_text))

    async def _l1_parse(self, state: AgentState, **_):
        """L1 — Parse raw text into structured profile JSON."""
        chunks = chunk_profile(state.raw_profile_text)
        if self.profile_parser:
            method = SafeMethodResolver.resolve(self.profile_parser, "parse")
            merged: dict = {}
            for chunk in chunks:
                partial = await _safe_call(method, chunk)
                if isinstance(partial, dict):
                    _deep_merge(merged, partial)
            state.extracted_profile = merged
        else:
            # Stub — replace with real LLM call
            state.extracted_profile = _stub_parse(state.raw_profile_text)

        _validate_profile(state.extracted_profile)
        profile_path = LOG_DIR / f"profile_{state.run_id}.json"
        profile_path.write_text(json.dumps(state.extracted_profile, indent=2))
        log.info("Profile written → %s", profile_path)
        state.mark_ok(1, skills_found=len(state.extracted_profile.get("skills", [])))

    async def _l2_plan(self, state: AgentState, job_targets=None, **_):
        """L2 — Build intent plan (target roles, salary, geo prefs)."""
        if self.intent_planner:
            method = SafeMethodResolver.resolve(self.intent_planner, "plan")
            state.intent_plan = await _safe_call(method, state.extracted_profile, job_targets)
        else:
            state.intent_plan = _stub_plan(state.extracted_profile, job_targets)
        state.mark_ok(2, roles=len(state.intent_plan.get("target_roles", [])))

    async def _l3_lead_scout(self, state: AgentState, **_):
        """L3 — Discover job leads; guard against timeouts and blocks."""
        if not self.lead_scout:
            log.warning("L3: LeadScout unavailable — using stub leads fallback.")
            state.job_leads = _stub_leads(state.intent_plan)
            state.mark_ok(3, leads_found=len(state.job_leads), fallback_mode="stub")
            return

        method = SafeMethodResolver.resolve(self.lead_scout, "search_jobs")
        try:
            leads = await asyncio.wait_for(
                _safe_call(method, state.intent_plan),
                timeout=120,   # 2-min hard cap per L3 timeout issue in checklist
            )
        except asyncio.TimeoutError:
            log.error("L3 Read Timeout — LeadScout exceeded 120s; continuing with 0 leads.")
            leads = []

        state.job_leads = leads if isinstance(leads, list) else []
        if not state.job_leads:
            log.warning("L3: discovery returned 0 leads — using stub leads fallback.")
            state.job_leads = _stub_leads(state.intent_plan)
            state.mark_ok(3, leads_found=len(state.job_leads), fallback_mode="stub")
            return
        state.mark_ok(3, leads_found=len(state.job_leads), fallback_mode="live")

    async def _l4_geo_fence(self, state: AgentState, **_):
        """L4 — Filter leads by location / remote preference."""
        if not state.job_leads:
            state.filtered_leads = []
            state.mark_ok(4, filtered=0)
            return

        if self.geo_fence:
            method = SafeMethodResolver.resolve(self.geo_fence, "filter_by_geo")
            result = await _safe_call(
                method,
                state.job_leads,
                state.intent_plan.get("geo_preferences", {}),
            )
            state.filtered_leads = result if isinstance(result, list) else state.job_leads
        else:
            state.filtered_leads = state.job_leads

        state.mark_ok(4, after_geo=len(state.filtered_leads))

    async def _l5_extract_score(self, state: AgentState, **_):
        """L5 — Extract JD requirements and score against profile.
           Threshold lowered to self.match_threshold (default 0.45).
        """
        if not state.filtered_leads:
            state.scored_jobs = []
            state.mark_ok(5, qualified=0)
            return

        if self.jd_extractor:
            method = SafeMethodResolver.resolve(self.jd_extractor, "extract_and_score")
            scored = await _safe_call(
                method,
                state.filtered_leads,
                state.extracted_profile,
                self.match_threshold,
            )
            state.scored_jobs = [j for j in (scored or []) if j.get("score", 0) >= self.match_threshold]
        else:
            state.scored_jobs = _stub_score(state.filtered_leads, self.match_threshold)

        state.mark_ok(5, qualified=len(state.scored_jobs), threshold=self.match_threshold)

    async def _l6_generate_artifacts(self, state: AgentState, **_):
        """L6 — Generate ATS resume + cover letter as real files in artifacts/."""
        if not state.scored_jobs:
            state.mark_ok(6, files_created=0)
            return

        if self.artifact_gen:
            method = SafeMethodResolver.resolve(self.artifact_gen, "generate_artifacts")
            artifacts = await _safe_call(
                method,
                state.extracted_profile,
                state.scored_jobs,
                ARTIFACTS_DIR,
            )
            state.artifacts = artifacts if isinstance(artifacts, dict) else {}
        else:
            state.artifacts = await _stub_generate_artifacts(
                state.extracted_profile, state.scored_jobs, ARTIFACTS_DIR
            )

        files_created = sum(len(v) for v in state.artifacts.values())
        state.mark_ok(6, files_created=files_created, artifact_dir=str(ARTIFACTS_DIR))

    async def _l7_apply(self, state: AgentState, **_):
        """L7 — ApplyExecutor: Playwright form-fill + file upload."""
        if not state.artifacts:
            log.warning("L7: No artifacts to submit — skipping ApplyExecutor.")
            state.layers[7].status = "skipped"
            return

        if not self.apply_executor:
            log.warning("L7: No ApplyExecutor injected — skipping.")
            state.layers[7].status = "skipped"
            return

        method = SafeMethodResolver.resolve(self.apply_executor, "apply_to_job")
        results = []
        for job in state.scored_jobs:
            job_id  = job.get("id", "unknown")
            job_url = job.get("url", "")
            files   = state.artifacts.get(job_id, {})

            if not job_url or not files:
                log.warning("L7: Skipping job_id=%s — missing URL or artifacts", job_id)
                continue

            try:
                result = await asyncio.wait_for(
                    _safe_call(method, job_url, files, job),
                    timeout=180,   # 3 min per application
                )
                results.append({"job_id": job_id, "status": "submitted", "detail": result})
                log.info("L7 ✓ Applied to job_id=%s  url=%s", job_id, job_url)
            except asyncio.TimeoutError:
                results.append({"job_id": job_id, "status": "timeout"})
                log.error("L7 ✗ Timeout on job_id=%s", job_id)
            except Exception as exc:
                results.append({"job_id": job_id, "status": "error", "error": str(exc)})
                log.error("L7 ✗ Error on job_id=%s: %s", job_id, exc)

        state.apply_results = results
        submitted = sum(1 for r in results if r["status"] == "submitted")
        state.mark_ok(7, submitted=submitted, attempted=len(results))

    async def _l8_confirm(self, state: AgentState, **_):
        """L8 — Capture confirmation numbers / screenshots."""
        confirmations = [r for r in state.apply_results if r.get("status") == "submitted"]
        state.mark_ok(8, confirmations=len(confirmations))

    async def _l9_flush(self, state: AgentState, **_):
        """L9 — Final state flush; UI progress bar reaches 100%."""
        final_report = {
            "run_id":        state.run_id,
            "progress":      state.progress_pct(),
            "layers":        [asdict(ls) for ls in state.layers],
            "leads_found":   len(state.job_leads),
            "qualified_jobs":len(state.scored_jobs),
            "applied_to":    sum(1 for r in state.apply_results if r.get("status") == "submitted"),
            "artifacts_dir": str(ARTIFACTS_DIR.resolve()),
            "errors":        state.errors,
        }
        report_path = LOG_DIR / f"report_{state.run_id}.json"
        report_path.write_text(json.dumps(final_report, indent=2))
        log.info("Final report → %s", report_path)
        state.mark_ok(9, report=str(report_path))


# ══════════════════════════════════════════════════════════════════════════════
# HELPER UTILITIES
# ══════════════════════════════════════════════════════════════════════════════

async def _safe_call(method: Callable, *args, **kwargs) -> Any:
    """Await coroutines; call plain functions directly."""
    result = method(*args, **kwargs)
    if inspect.isawaitable(result):
        result = await result
    return result


def _deep_merge(base: dict, update: dict) -> None:
    for k, v in update.items():
        if k in base and isinstance(base[k], list) and isinstance(v, list):
            # Merge lists (e.g. skills from multiple chunks)
            seen = {json.dumps(i, sort_keys=True) for i in base[k]}
            for item in v:
                if json.dumps(item, sort_keys=True) not in seen:
                    base[k].append(item)
        elif k in base and isinstance(base[k], dict) and isinstance(v, dict):
            _deep_merge(base[k], v)
        else:
            base[k] = v


def _validate_profile(profile: dict) -> None:
    required = ["name", "skills", "experience"]
    missing  = [f for f in required if not profile.get(f)]
    if missing:
        raise ValueError(f"Extracted profile missing required fields: {missing}")


# ══════════════════════════════════════════════════════════════════════════════
# STUB IMPLEMENTATIONS  (replace with real LLM / API calls)
# ══════════════════════════════════════════════════════════════════════════════

def _stub_parse(text: str) -> dict:
    log.warning("Using stub profile parser — replace with real LLM parser.")
    return {
        "name":       "Candidate",
        "email":      "",
        "phone":      "",
        "skills":     ["Python", "SQL", "Communication"],
        "experience": [{"title": "Software Engineer", "years": 3}],
        "education":  [],
        "summary":    text[:300],
    }


def _stub_plan(profile: dict, job_targets: Optional[list[str]]) -> dict:
    return {
        "target_roles":    job_targets or ["Software Engineer", "Backend Developer"],
        "salary_min_usd":  100_000,
        "salary_max_usd":  160_000,
        "geo_preferences": {"remote": True, "locations": []},
        "keywords":        profile.get("skills", [])[:10],
    }


def _stub_leads(intent_plan: dict) -> list[dict]:
    roles = intent_plan.get("target_roles") or ["Software Engineer"]
    role = str(roles[0])
    return [
        {
            "id": "stub_001",
            "title": f"Senior {role}",
            "company": "DemoTech",
            "url": "https://boards.greenhouse.io/demotech/jobs/123",
            "location": "Remote",
            "remote": True,
            "description": "Production APIs using Python, SQL, and backend communication in cloud systems.",
            "source": "stub",
        },
        {
            "id": "stub_002",
            "title": role,
            "company": "NextGen Labs",
            "url": "https://jobs.lever.co/nextgen/456",
            "location": "United States",
            "remote": True,
            "description": "Build scalable backend with Python, SQL, FastAPI, and communication-heavy workflows.",
            "source": "stub",
        },
    ]


def _stub_score(leads: list[dict], threshold: float) -> list[dict]:
    import random
    scored = []
    for lead in leads:
        score = random.uniform(0.3, 0.95)
        if score >= threshold:
            scored.append({**lead, "score": round(score, 3)})
    return scored


async def _stub_generate_artifacts(
    profile: dict, jobs: list[dict], out_dir: Path
) -> dict:
    """
    Generates real .docx and .pdf stubs using python-docx + reportlab.
    Replace with your LLM-powered template renderer.
    """
    artifacts = {}
    for job in jobs:
        job_id = job.get("id", f"job_{id(job)}")
        job_dir = out_dir / job_id
        job_dir.mkdir(parents=True, exist_ok=True)

        resume_path = job_dir / "resume.docx"
        cover_path  = job_dir / "cover_letter.docx"

        _write_docx_resume(profile, job, resume_path)
        _write_docx_cover(profile, job, cover_path)

        # Convert to PDF if available
        resume_pdf = _docx_to_pdf(resume_path)
        cover_pdf  = _docx_to_pdf(cover_path)

        artifacts[job_id] = {
            "resume_docx": str(resume_path),
            "cover_docx":  str(cover_path),
            "resume_pdf":  str(resume_pdf) if resume_pdf else None,
            "cover_pdf":   str(cover_pdf)  if cover_pdf  else None,
        }
        log.info("L6 Artifacts created for job_id=%s → %s", job_id, job_dir)
    return artifacts


def _write_docx_resume(profile: dict, job: dict, path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading(profile.get("name", "Candidate"), 0)
        doc.add_paragraph(f"Email: {profile.get('email', '')}  |  Phone: {profile.get('phone', '')}")
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(profile.get("summary", ""))
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(profile.get("skills", [])))
        doc.add_heading("Experience", level=1)
        for exp in profile.get("experience", []):
            p = doc.add_paragraph()
            r = p.add_run(exp.get("title", ""))
            r.bold = True
            doc.add_paragraph(f"  {exp.get('years', '')} years")
        doc.save(path)
    except ImportError:
        path.write_text(f"RESUME PLACEHOLDER\n{json.dumps(profile, indent=2)}")
        log.warning("python-docx not installed — wrote text placeholder for resume.")


def _write_docx_cover(profile: dict, job: dict, path: Path) -> None:
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Cover Letter", 0)
        doc.add_paragraph(f"Applying for: {job.get('title', 'the position')}")
        doc.add_paragraph(f"Company: {job.get('company', '')}")
        doc.add_paragraph(
            f"Dear Hiring Manager,\n\n"
            f"I am excited to apply for the {job.get('title', 'open position')} role at "
            f"{job.get('company', 'your company')}. With expertise in "
            f"{', '.join(profile.get('skills', [])[:3])}, I am confident I can contribute "
            f"meaningfully to your team.\n\nSincerely,\n{profile.get('name', 'Candidate')}"
        )
        doc.save(path)
    except ImportError:
        path.write_text(f"COVER LETTER PLACEHOLDER\n{job.get('title', '')}")
        log.warning("python-docx not installed — wrote text placeholder for cover letter.")


def _docx_to_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        import subprocess
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(docx_path.parent), str(docx_path)],
            capture_output=True, timeout=30
        )
        if result.returncode == 0 and pdf_path.exists():
            return pdf_path
    except Exception as exc:
        log.debug("PDF conversion skipped (%s)", exc)
    return None


# ══════════════════════════════════════════════════════════════════════════════
# APPLY EXECUTOR  (L7 — Playwright form-fill + upload)
# ══════════════════════════════════════════════════════════════════════════════

class ApplyExecutor:
    """
    Playwright-based job application executor.

    Strategy:
      1. Navigate to job URL
      2. Detect application form (iframe or redirect)
      3. Fill standard fields (name, email, phone, LinkedIn)
      4. Upload resume PDF (preferred) or DOCX fallback
      5. Upload cover letter if upload slot exists
      6. Submit form
      7. Capture confirmation text / screenshot

    ATS Detection Map covers: Greenhouse, Lever, Workday, iCIMS, Taleo,
    BambooHR, SmartRecruiters, and generic HTML forms.
    """

    ATS_PATTERNS = {
        "greenhouse":     "boards.greenhouse.io",
        "lever":          "jobs.lever.co",
        "workday":        "myworkdayjobs.com",
        "icims":          "icims.com",
        "taleo":          "taleo.net",
        "bamboohr":       "bamboohr.com",
        "smartrecruiters":"smartrecruiters.com",
    }

    UPLOAD_SELECTORS = [
        # Ordered from most → least specific
        "input[type='file'][name*='resume']",
        "input[type='file'][name*='cv']",
        "input[type='file'][accept*='pdf']",
        "input[type='file'][accept*='docx']",
        "input[type='file'][accept*='doc']",
        "input[type='file']",
        "[data-testid*='upload']",
        "[aria-label*='upload' i]",
        "button:has-text('Upload')",
        "label:has-text('Resume')",
        "label:has-text('CV')",
    ]

    COVER_SELECTORS = [
        "input[type='file'][name*='cover']",
        "input[type='file'][name*='letter']",
        "[data-testid*='cover']",
        "[aria-label*='cover' i]",
    ]

    SUBMIT_SELECTORS = [
        "button[type='submit']",
        "input[type='submit']",
        "button:has-text('Submit')",
        "button:has-text('Apply')",
        "button:has-text('Send Application')",
        "[data-testid*='submit']",
        "[aria-label*='submit' i]",
    ]

    async def apply_to_job(
        self,
        job_url: str,
        files: dict,
        job_meta: dict,
        profile: Optional[dict] = None,
    ) -> dict:
        """
        Main entry point.  files = {
            "resume_pdf": "/path/to/resume.pdf",
            "cover_pdf":  "/path/to/cover.pdf",
            "resume_docx": ...,
            "cover_docx": ...,
        }
        """
        try:
            from playwright.async_api import async_playwright, TimeoutError as PWTimeout
        except ImportError:
            raise RuntimeError(
                "Playwright not installed. Run: pip install playwright && playwright install chromium"
            )

        resume_file = files.get("resume_pdf") or files.get("resume_docx")
        cover_file  = files.get("cover_pdf")  or files.get("cover_docx")

        if not resume_file or not Path(resume_file).exists():
            raise FileNotFoundError(f"Resume file not found: {resume_file}")

        async with async_playwright() as pw:
            browser = await pw.chromium.launch(headless=True, args=["--no-sandbox"])
            ctx     = await browser.new_context(
                user_agent=(
                    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                    "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
                ),
                accept_downloads=True,
            )
            page = await ctx.new_page()

            result = {
                "url":          job_url,
                "ats":          self._detect_ats(job_url),
                "resume_uploaded": False,
                "cover_uploaded":  False,
                "submitted":    False,
                "confirmation": None,
                "screenshot":   None,
            }

            try:
                await page.goto(job_url, wait_until="domcontentloaded", timeout=30_000)
                await page.wait_for_timeout(2000)   # let JS hydrate

                # ── Fill text fields ──────────────────────────────────────
                if profile:
                    await self._fill_text_fields(page, profile)

                # ── Upload Resume ─────────────────────────────────────────
                uploaded = await self._upload_file(
                    page, self.UPLOAD_SELECTORS, resume_file, "resume"
                )
                result["resume_uploaded"] = uploaded

                # ── Upload Cover Letter (optional) ────────────────────────
                if cover_file and Path(cover_file).exists():
                    cover_uploaded = await self._upload_file(
                        page, self.COVER_SELECTORS, cover_file, "cover letter"
                    )
                    result["cover_uploaded"] = cover_uploaded

                # ── Submit ────────────────────────────────────────────────
                submitted = await self._click_submit(page)
                result["submitted"] = submitted

                if submitted:
                    await page.wait_for_timeout(3000)
                    result["confirmation"] = await self._capture_confirmation(page)

            except PWTimeout as exc:
                log.error("Playwright timeout on %s: %s", job_url, exc)
                result["error"] = f"Timeout: {exc}"
            finally:
                # Always capture screenshot for debugging
                ss_path = ARTIFACTS_DIR / f"screenshot_{datetime.now().strftime('%H%M%S')}.png"
                await page.screenshot(path=str(ss_path), full_page=True)
                result["screenshot"] = str(ss_path)
                await browser.close()

        return result

    def _detect_ats(self, url: str) -> str:
        for name, pattern in self.ATS_PATTERNS.items():
            if pattern in url:
                return name
        return "generic"

    async def _fill_text_fields(self, page, profile: dict) -> None:
        """Best-effort autofill of standard form fields."""
        field_map = {
            # Selector patterns → profile key
            "input[name*='first' i], input[placeholder*='first' i]":
                profile.get("name", "").split()[0] if profile.get("name") else "",
            "input[name*='last' i], input[placeholder*='last' i]":
                " ".join(profile.get("name", "").split()[1:]),
            "input[type='email'], input[name*='email' i]":
                profile.get("email", ""),
            "input[type='tel'], input[name*='phone' i]":
                profile.get("phone", ""),
            "input[name*='linkedin' i], input[placeholder*='linkedin' i]":
                profile.get("linkedin_url", ""),
        }
        for selector, value in field_map.items():
            if not value:
                continue
            try:
                el = page.locator(selector).first
                if await el.count() > 0:
                    await el.fill(value)
            except Exception:
                pass   # Non-fatal — continue filling remaining fields

    async def _upload_file(
        self, page, selectors: list[str], file_path: str, label: str
    ) -> bool:
        """Try each selector in order; return True if file was set."""
        for selector in selectors:
            try:
                locator = page.locator(selector).first
                if await locator.count() == 0:
                    continue
                tag = await locator.evaluate("el => el.tagName.toLowerCase()")
                if tag == "input":
                    await locator.set_input_files(file_path)
                    log.info("L7 ✓ Uploaded %s via selector '%s'", label, selector)
                    return True
                else:
                    # Visible label/button — click to open native file dialog
                    async with page.expect_file_chooser() as fc_info:
                        await locator.click()
                    fc = await fc_info.value
                    await fc.set_files(file_path)
                    log.info("L7 ✓ Uploaded %s via file-chooser '%s'", label, selector)
                    return True
            except Exception as exc:
                log.debug("_upload_file: selector '%s' failed: %s", selector, exc)
                continue
        log.warning("L7 ✗ Could not find upload slot for %s", label)
        return False

    async def _click_submit(self, page) -> bool:
        for selector in self.SUBMIT_SELECTORS:
            try:
                btn = page.locator(selector).first
                if await btn.count() > 0 and await btn.is_enabled():
                    await btn.click()
                    log.info("L7 ✓ Clicked submit via '%s'", selector)
                    return True
            except Exception:
                continue
        log.warning("L7 ✗ Submit button not found")
        return False

    async def _capture_confirmation(self, page) -> Optional[str]:
        """Extract confirmation text from common thank-you / confirmation patterns."""
        patterns = [
            "[class*='confirm' i]", "[class*='success' i]", "[class*='thank' i]",
            "h1", "h2", "[role='alert']", "[aria-live]",
        ]
        for sel in patterns:
            try:
                el = page.locator(sel).first
                if await el.count() > 0:
                    text = (await el.inner_text()).strip()
                    if text:
                        return text[:500]
            except Exception:
                continue
        return None


# ══════════════════════════════════════════════════════════════════════════════
# WIRING EXAMPLE  (replace stubs with real service instances)
# ══════════════════════════════════════════════════════════════════════════════

async def main():
    """Quick smoke-test — swap stubs for production services."""
    sample_profile = """
    John Smith | john@example.com | +1-555-0100
    Senior Software Engineer with 6 years of experience building scalable backend
    systems in Python, FastAPI, and PostgreSQL. Led a team of 4 engineers at Acme Corp.
    Strong background in cloud infrastructure (AWS, GCP), CI/CD (GitHub Actions),
    and distributed systems. Open to remote or hybrid roles in the US.
    Education: B.Sc. Computer Science, MIT, 2017.
    """

    orchestrator = Orchestrator(
        apply_executor=ApplyExecutor(),
        match_threshold=0.45,
    )
    state = await orchestrator.run(sample_profile, job_targets=["Senior Backend Engineer"])

    print("\n━━━ Run Complete ━━━")
    print(f"Progress : {state.progress_pct()}%")
    print(f"Leads    : {len(state.job_leads)}")
    print(f"Qualified: {len(state.scored_jobs)}")
    print(f"Applied  : {sum(1 for r in state.apply_results if r.get('status')=='submitted')}")
    print(f"Errors   : {state.errors or 'none'}")
    print(f"Artifacts: {ARTIFACTS_DIR.resolve()}")


if __name__ == "__main__":
    asyncio.run(main())
