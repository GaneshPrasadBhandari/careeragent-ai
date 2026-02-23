# src/careeragent/langgraph/hitl_flows.py
from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any, Dict, List


def _write_docx_from_md(path: Path, md: str) -> None:
    """Description: Convert markdown to simple ATS DOCX.
    Layer: L6
    Input: md
    Output: saved file
    """
    try:
        import docx  # type: ignore
        doc = docx.Document()
        for ln in (md or "").splitlines():
            ln = ln.rstrip()
            if not ln.strip():
                continue
            if ln.startswith("# "):
                doc.add_heading(ln[2:].strip(), level=0)
            elif ln.startswith("## "):
                doc.add_heading(ln[3:].strip(), level=1)
            elif ln.lstrip().startswith("- "):
                doc.add_paragraph(ln.lstrip()[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(ln.strip())
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
    except Exception:
        return


def _write_pdf_from_md(path: Path, md: str) -> None:
    """Description: Convert markdown to simple single-column PDF.
    Layer: L6
    Input: md
    Output: saved file
    """
    try:
        from reportlab.lib.pagesizes import LETTER  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
        path.parent.mkdir(parents=True, exist_ok=True)
        c = canvas.Canvas(str(path), pagesize=LETTER)
        w, h = LETTER
        x, y = 50, h - 50
        lh = 12

        def draw(text: str) -> None:
            nonlocal y
            c.setFont("Helvetica", 10)
            c.drawString(x, y, text)
            y -= lh

        for ln in (md or "").splitlines():
            if y < 60:
                c.showPage()
                y = h - 50
            ln = ln.rstrip()
            if not ln.strip():
                y -= lh
                continue
            if ln.startswith("# "):
                c.setFont("Helvetica-Bold", 16)
                c.drawString(x, y, ln[2:].strip())
                y -= 18
                continue
            if ln.startswith("## "):
                c.setFont("Helvetica-Bold", 12)
                c.drawString(x, y, ln[3:].strip())
                y -= 14
                continue
            if ln.lstrip().startswith("- "):
                draw("• " + ln.lstrip()[2:].strip())
                continue
            draw(ln.strip())
        c.save()
    except Exception:
        return


def _artifacts_root() -> Path:
    return Path("src/careeragent/artifacts").resolve()


def _runs_dir(run_id: str) -> Path:
    d = _artifacts_root() / "runs" / run_id
    d.mkdir(parents=True, exist_ok=True)
    return d


def _write_text(path: Path, text: str) -> str:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")
    return str(path)


def _job_key(url: str) -> str:
    return hashlib.md5(url.encode("utf-8"), usedforsecurity=False).hexdigest()[:12]  # noqa: S324


def _fallback_resume_md(profile: Dict[str, Any], job: Dict[str, Any]) -> str:
    name = (profile.get("name") or "Candidate").strip()
    contact = profile.get("contact") or {}
    skills = profile.get("skills") or []
    title = job.get("title") or "Target Role"
    matched = job.get("matched_skills") or []
    missing = job.get("missing_skills") or []

    return f"""# {name}
{contact.get('phone','')} | {contact.get('email','')}
{contact.get('linkedin','')}

## Target Role
**{title}**

## Professional Summary
AI/ML + GenAI engineer focused on production-grade delivery (MLOps, CI/CD, evaluation, guardrails). ATS-friendly and metric-driven.

## Core Skills (ATS)
{", ".join(skills[:30])}

## Matched Skills
{", ".join(matched[:20]) if matched else "See Core Skills"}

## Skill Gaps
{", ".join(missing[:15]) if missing else "None detected from job text"}

## Experience
- Add 4–6 bullets per role with metrics (latency, cost, accuracy, scale, adoption)
- Include tools + scope + impact

## Education
- MSIT (Healthcare Tech) — Clark University (in progress)

## Projects
- CareerAgent-AI: LangGraph orchestration + HITL + explainability + analytics
"""


def _fallback_cover_md(profile: Dict[str, Any], job: Dict[str, Any], country: str) -> str:
    name = (profile.get("name") or "Candidate").strip()
    contact = profile.get("contact") or {}
    title = job.get("title") or "the role"
    url = job.get("url") or job.get("job_id") or ""

    greeting = "Dear Hiring Manager," if country.upper() == "US" else "Dear Hiring Team,"
    return f"""# Cover Letter — {name}

{contact.get('email','')} | {contact.get('phone','')} | {contact.get('linkedin','')}

{greeting}

I’m applying for **{title}**. I build production-grade AI/ML and GenAI systems with strong MLOps discipline (tracking, reproducibility, governance).
I focus on measurable outcomes and reliability: evaluation, guardrails, and automation.

Highlights:
- End-to-end AI delivery: data → training → evaluation → deployment
- GenAI/LLM apps: RAG, orchestration, tool selection, safety gates
- Cloud + DevOps: CI/CD, containers, monitoring, artifact/version control

Sincerely,  
**{name}**

Job link: {url}
"""


async def approve_ranking_flow(state: Dict[str, Any]) -> Dict[str, Any]:
    run_id = str(state.get("run_id") or "run")
    prefs = state.get("preferences") or {}
    country = str(prefs.get("country") or "US")

    ranking: List[Dict[str, Any]] = state.get("ranking") or []
    approved: List[str] = (state.get("meta") or {}).get("approved_job_urls") or []
    approved = [u.strip() for u in approved if isinstance(u, str) and u.strip()]

    if approved:
        approved_set = set(approved)
        ranking = [j for j in ranking if str(j.get("url") or j.get("job_id") or "") in approved_set]

    if not ranking:
        ranking = (state.get("ranking") or [])[:10]

    # Preferred: real L6 generation
    try:
        from careeragent.langgraph.nodes_l6_l9 import l6_draft_node, l6_evaluator_node
        state["ranking"] = ranking
        state.update(await l6_draft_node(state))      # type: ignore[arg-type]
        state.update(await l6_evaluator_node(state))  # type: ignore[arg-type]
    except Exception:
        prof = state.get("profile") or {}
        run_dir = _runs_dir(run_id)
        artifacts = state.setdefault("artifacts", {})
        drafts = []

        for job in ranking[:10]:
            url = str(job.get("url") or job.get("job_id") or "")
            key = _job_key(url or (job.get("title") or "job"))
            resume_p = run_dir / f"resume_{key}.md"
            cover_p = run_dir / f"cover_{key}.md"

            resume_md = _fallback_resume_md(prof, job)
            cover_md = _fallback_cover_md(prof, job, country)
            artifacts[f"resume_{key}"] = {"path": _write_text(resume_p, resume_md), "content_type": "text/markdown"}
            artifacts[f"cover_{key}"] = {"path": _write_text(cover_p, cover_md), "content_type": "text/markdown"}

            # Also export DOCX/PDF for demo
            resume_docx = run_dir / f"resume_{key}.docx"
            resume_pdf = run_dir / f"resume_{key}.pdf"
            cover_docx = run_dir / f"cover_{key}.docx"
            cover_pdf = run_dir / f"cover_{key}.pdf"
            _write_docx_from_md(resume_docx, resume_md)
            _write_pdf_from_md(resume_pdf, resume_md)
            _write_docx_from_md(cover_docx, cover_md)
            _write_pdf_from_md(cover_pdf, cover_md)

            if resume_docx.exists():
                artifacts[f"resume_docx_{key}"] = {"path": str(resume_docx), "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
            if resume_pdf.exists():
                artifacts[f"resume_pdf_{key}"] = {"path": str(resume_pdf), "content_type": "application/pdf"}
            if cover_docx.exists():
                artifacts[f"cover_docx_{key}"] = {"path": str(cover_docx), "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
            if cover_pdf.exists():
                artifacts[f"cover_pdf_{key}"] = {"path": str(cover_pdf), "content_type": "application/pdf"}

            drafts.append({"job_url": url, "job_title": job.get("title"), "resume_path": artifacts[f"resume_{key}"]["path"], "cover_path": artifacts[f"cover_{key}"]["path"]})

        state["drafts"] = {"drafts": drafts}

    state.setdefault("live_feed", []).append({"layer": "L6", "agent": "HITL", "message": f"Drafts generated for {min(10, len(ranking))} jobs. Review drafts."})
    state["status"] = "needs_human_approval"
    state["pending_action"] = "review_drafts"
    return state


async def approve_drafts_flow(state: Dict[str, Any]) -> Dict[str, Any]:
    try:
        from careeragent.langgraph.nodes_l6_l9 import (
            l7_apply_node, l7_evaluator_node,
            l8_tracker_node, l8_evaluator_node,
            l9_analytics_node,
        )
        state.update(await l7_apply_node(state))       # type: ignore[arg-type]
        state.update(await l7_evaluator_node(state))   # type: ignore[arg-type]
        if state.get("status") == "needs_human_approval":
            return state

        state.update(await l8_tracker_node(state))     # type: ignore[arg-type]
        state.update(await l8_evaluator_node(state))   # type: ignore[arg-type]
        if state.get("status") == "needs_human_approval":
            return state

        state.update(await l9_analytics_node(state))   # type: ignore[arg-type]
    except Exception:
        pass

    state.setdefault("live_feed", []).append({"layer": "L9", "agent": "HITL", "message": "Run completed."})
    state["status"] = "completed"
    state["pending_action"] = None
    return state