"""drafting_agent_service.py — L6 ATS Resume + Cover Letter Generator (Fixed).

Key fixes:
  - Proper ATS resume section ordering: # Name → Contact → ## Summary → ## Skills → ## Experience → ## Education
  - Professional PDF via reportlab Platypus (not canvas.drawString plain text)
  - Gemini token budget 4000 (was 1500) for full resume generation
  - Profile sent as clean JSON string (not Python repr)
  - Cover letter uses real experience bullets for Action→Result→Challenge
  - All formats: .md + .docx + .pdf for resume and cover letter
"""
from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from careeragent.core.settings import Settings
from careeragent.core.state import AgentState, ArtifactRef
from careeragent.nlp.skills import compute_jd_alignment, extract_skills, normalize_skill
from careeragent.tools.llm_tools import GeminiClient
from careeragent.tools.web_tools import stable_key


# ─────────────────────────────────────────────
# HTML escape helper
# ─────────────────────────────────────────────
def _esc(text: str) -> str:
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _md_to_html(text: str) -> str:
    text = _esc(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    return text


# ─────────────────────────────────────────────
# DOCX writer — ATS-safe
# ─────────────────────────────────────────────
def _write_docx(text: str, path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if stripped.startswith("# "):
                h = doc.add_heading(stripped[2:].strip(), level=0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in h.runs:
                    run.font.size = Pt(16)
                    run.font.bold = True
            elif stripped.startswith("## "):
                h = doc.add_heading(stripped[3:].strip().upper(), level=1)
                for run in h.runs:
                    run.font.size = Pt(11)
                    run.font.bold = True
            elif stripped.startswith("### "):
                h = doc.add_heading(stripped[4:].strip(), level=2)
                for run in h.runs:
                    run.font.size = Pt(10.5)
                    run.font.bold = True
            elif re.match(r"^[-•·*]\s+", stripped):
                p = doc.add_paragraph(style="List Bullet")
                p.text = re.sub(r"^[-•·*]\s+", "", stripped)
                for run in p.runs:
                    run.font.size = Pt(10)
            else:
                p = doc.add_paragraph(stripped)
                for run in p.runs:
                    run.font.size = Pt(10)

        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
    except Exception:
        pass


# ─────────────────────────────────────────────
# PDF writer — professional ATS resume layout
# ─────────────────────────────────────────────
def _write_pdf(text: str, path: Path) -> None:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        path.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(
            str(path), pagesize=LETTER,
            rightMargin=0.75 * inch, leftMargin=0.75 * inch,
            topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        )

        style_name = ParagraphStyle(
            "Name", fontSize=17, fontName="Helvetica-Bold",
            spaceAfter=4, alignment=TA_CENTER, textColor=colors.HexColor("#1a1a1a"),
        )
        style_contact = ParagraphStyle(
            "Contact", fontSize=9, fontName="Helvetica",
            spaceAfter=6, alignment=TA_CENTER, textColor=colors.HexColor("#444444"),
        )
        style_section = ParagraphStyle(
            "Section", fontSize=11, fontName="Helvetica-Bold",
            spaceBefore=8, spaceAfter=2, textColor=colors.HexColor("#1a1a1a"),
        )
        style_role = ParagraphStyle(
            "Role", fontSize=10, fontName="Helvetica-Bold", spaceAfter=1, spaceBefore=6,
        )
        style_body = ParagraphStyle(
            "Body", fontSize=9.5, fontName="Helvetica", spaceAfter=2, leading=13,
        )
        style_bullet = ParagraphStyle(
            "Bullet", fontSize=9.5, fontName="Helvetica",
            leftIndent=14, spaceAfter=1, leading=13, bulletIndent=4,
        )

        story = []
        lines = text.splitlines()
        is_header_done = False

        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 4))
                continue

            if stripped.startswith("# "):
                story.append(Paragraph(_esc(stripped[2:].strip()), style_name))
                is_header_done = False

            elif stripped.startswith("## "):
                story.append(Spacer(1, 4))
                story.append(HRFlowable(width="100%", thickness=0.7, color=colors.HexColor("#333333")))
                story.append(Paragraph(_esc(stripped[3:].strip().upper()), style_section))

            elif stripped.startswith("### "):
                role_text = stripped[4:].strip()
                parts = re.split(r"\s+[|—–-]\s+", role_text, maxsplit=1)
                if len(parts) == 2:
                    html = f"<b>{_esc(parts[0])}</b>&nbsp;|&nbsp;{_esc(parts[1])}"
                else:
                    html = f"<b>{_esc(role_text)}</b>"
                story.append(Paragraph(html, style_role))

            elif re.match(r"^[-•·*]\s+", stripped):
                bullet_text = re.sub(r"^[-•·*]\s+", "", stripped)
                story.append(Paragraph(f"• {_md_to_html(bullet_text)}", style_bullet))

            else:
                html_line = _md_to_html(stripped)
                s = style_contact if (not is_header_done and ("@" in stripped or "linkedin" in stripped.lower() or "|" in stripped)) else style_body
                story.append(Paragraph(html_line, s))

        doc.build(story)

    except Exception:
        # Last-resort plain PDF
        try:
            from reportlab.pdfgen import canvas as rlcanvas
            from reportlab.lib.pagesizes import LETTER
            path.parent.mkdir(parents=True, exist_ok=True)
            c = rlcanvas.Canvas(str(path), pagesize=LETTER)
            w, h = LETTER
            y = h - 50
            for line in text.splitlines():
                if y < 60:
                    c.showPage(); y = h - 50
                c.setFont("Helvetica", 9)
                c.drawString(50, y, line[:120])
                y -= 13
            c.save()
        except Exception:
            pass


# ─────────────────────────────────────────────
# ATS scoring
# ─────────────────────────────────────────────
def ats_format_check_percent(resume_md: str) -> float:
    txt = resume_md or ""
    score = 100.0
    if "|" in txt and re.search(r"\n\s*\|", txt):
        score -= 35
    if not re.search(r"^#\s+", txt, flags=re.M):
        score -= 20
    if not re.search(r"^##\s+", txt, flags=re.M):
        score -= 15
    bullets = len(re.findall(r"^\s*[-•]\s+", txt, flags=re.M))
    if bullets < 5:
        score -= 15
    for sec in ["experience", "education", "skills"]:
        if not re.search(sec, txt, flags=re.I):
            score -= 5
    return float(max(0.0, min(100.0, score)))


def ats_keyword_match_percent(job_text: str, resume_text: str, *, profile_skills: List[str]) -> float:
    prof = [normalize_skill(s) for s in (profile_skills or []) if s]
    jd_skills = set(extract_skills(job_text, extra_candidates=prof))
    resume_skills = set(extract_skills(resume_text, extra_candidates=prof))
    if not jd_skills:
        return 0.0
    matched = len(jd_skills & resume_skills)
    return round((matched / max(1, len(jd_skills))) * 100.0, 2)


def optimize_resume_keywords(*, resume_md: str, jd_text: str, profile_skills: List[str], top_n: int = 5) -> Tuple[str, List[str]]:
    """Reverse-ATS pass: inject top missing JD keywords into Skills Highlights section."""
    jd_terms = [s for s in extract_skills(jd_text, extra_candidates=profile_skills) if s]
    resume_terms = set(extract_skills(resume_md, extra_candidates=profile_skills))
    missing: List[str] = []
    for term in jd_terms:
        if term in resume_terms or term in missing:
            continue
        missing.append(term)
        if len(missing) >= top_n:
            break
    if not missing:
        return resume_md, []

    section_header = "## Skills Highlights"
    section_body = "\n".join([f"- {m}" for m in missing])
    if section_header in resume_md:
        enriched = resume_md.replace(section_header, f"{section_header}\n{section_body}", 1)
    else:
        enriched = f"{resume_md.strip()}\n\n{section_header}\n{section_body}\n"
    return enriched, missing


def _split_words(text: str) -> List[str]:
    return [w.lower() for w in re.findall(r"[a-zA-Z][a-zA-Z0-9+/.-]{2,}", str(text or ""))]


def _top_relevant_projects(profile: Dict[str, Any], jd_text: str, *, limit: int = 4) -> List[Dict[str, Any]]:
    """Rank past experience entries against JD keywords for contextual synthesis."""
    jd_terms = set(_split_words(jd_text))
    ranked: List[Tuple[int, Dict[str, Any]]] = []
    for e in profile.get("experience") or []:
        haystack = " ".join(
            [
                str(e.get("title") or ""),
                str(e.get("company") or ""),
                " ".join([str(b) for b in (e.get("bullets") or [])]),
            ]
        )
        tokens = _split_words(haystack)
        overlap = len(jd_terms.intersection(tokens))
        ranked.append((overlap, e))
    ranked.sort(key=lambda item: item[0], reverse=True)
    return [item[1] for item in ranked[: max(1, limit)] if item[1]]


class DraftingAgentService:
    """Description: Generate ATS-friendly resume + US-style story cover letter.
    Layer: L6
    Input: approved jobs + extracted_profile
    Output: resume_*, cover_* artifacts (MD + DOCX + PDF)
    """

    def __init__(self, settings: Settings) -> None:
        self.s = settings
        self.gemini = GeminiClient(settings, model="gemini-1.5-pro")

    def generate_for_jobs(self, state: AgentState) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        prof = state.extracted_profile
        profile_skills = list(prof.get("skills") or [])

        approved = list(state.approved_job_urls[: state.preferences.draft_count])
        if not approved:
            # Fallback: use top ranked jobs
            approved = [
                j.get("url") for j in state.ranking[: state.preferences.draft_count]
                if j.get("url")
            ]

        run_dir = Path("src/careeragent/artifacts/runs") / state.run_id
        run_dir.mkdir(parents=True, exist_ok=True)

        for url in approved:
            job = next((j for j in state.ranking if str(j.get("url") or "") == url), None)
            if not job:
                job = next((j for j in state.jobs_scored if str(j.get("url") or "") == url), None)
            if not job:
                continue

            key = stable_key(url)
            jd = job.get("full_text_md") or job.get("full_text") or job.get("snippet") or ""
            title = re.sub(r"\s*[|-].*$", "", job.get("title") or "Target Role").strip()[:80]

            revision_notes = str((state.meta or {}).get("draft_revision_feedback") or "")
            resume_md, cover_md = self._gen_with_llm(prof, jd, title)
            resume_md, injected = optimize_resume_keywords(resume_md=resume_md, jd_text=jd, profile_skills=profile_skills, top_n=5)
            resume_plain = _strip_markdown(resume_md)
            cover_plain = _strip_markdown(cover_md)

            # Scoring
            layout_pct = ats_format_check_percent(resume_md)
            align = compute_jd_alignment(
                jd_text=jd,
                resume_skills=extract_skills(resume_md, extra_candidates=profile_skills),
            )
            kw_pct = ats_keyword_match_percent(jd, resume_md, profile_skills=profile_skills)

            job["ats_layout_percent"] = layout_pct
            job["jd_alignment_percent"] = align.jd_alignment_percent
            job["missing_skills_gap_percent"] = align.missing_skills_gap_percent
            job["ats_keyword_match_percent"] = kw_pct
            job["missing_jd_skills"] = align.missing_jd_skills[:40]
            job["matched_jd_skills"] = align.matched_jd_skills[:40]
            job["ats_keyword_injected"] = injected

            # Write all formats
            resume_md_path = run_dir / f"resume_{key}.md"
            cover_md_path = run_dir / f"cover_{key}.md"
            resume_docx = run_dir / f"resume_{key}.docx"
            cover_docx = run_dir / f"cover_{key}.docx"
            resume_pdf = run_dir / f"resume_{key}.pdf"
            cover_pdf = run_dir / f"cover_{key}.pdf"

            resume_md_path.write_text(resume_md, encoding="utf-8")
            cover_md_path.write_text(cover_md, encoding="utf-8")
            _write_docx(resume_md, resume_docx)
            _write_docx(cover_md, cover_docx)
            _write_pdf(resume_md, resume_pdf)
            _write_pdf(cover_md, cover_pdf)

            # Register artifacts
            state.artifacts[f"resume_{key}"] = ArtifactRef(path=str(resume_md_path), mime="text/markdown")
            state.artifacts[f"cover_{key}"] = ArtifactRef(path=str(cover_md_path), mime="text/markdown")
            state.artifacts[f"resume_{key}_docx"] = ArtifactRef(
                path=str(resume_docx),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            state.artifacts[f"cover_{key}_docx"] = ArtifactRef(
                path=str(cover_docx),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            state.artifacts[f"resume_{key}_pdf"] = ArtifactRef(path=str(resume_pdf), mime="application/pdf")
            state.artifacts[f"cover_{key}_pdf"] = ArtifactRef(path=str(cover_pdf), mime="application/pdf")

            out.append({
                "job_url": url,
                "job_title": title,
                "key": key,
                "ats_layout_percent": layout_pct,
                "jd_alignment_percent": align.jd_alignment_percent,
                "missing_skills_gap_percent": align.missing_skills_gap_percent,
                "ats_keyword_match_percent": kw_pct,
                "resume_md_path": str(resume_md_path),
                "cover_md_path": str(cover_md_path),
                "resume_pdf_path": str(resume_pdf),
                "resume_docx_path": str(resume_docx),
            })

        return out

    def _gen_with_llm(self, profile: Dict[str, Any], jd: str, title: str) -> Tuple[str, str]:
        fallback_resume = _build_fallback_resume(profile, title, jd)
        fallback_cover = _build_fallback_cover(profile, title)

        if not self.s.GEMINI_API_KEY:
            return fallback_resume, fallback_cover

        # Build clean structured profile JSON
        clean_profile = {
            "name": profile.get("name"),
            "contact": profile.get("contact") or {},
            "summary": profile.get("summary"),
            "skills": (profile.get("skills") or [])[:60],
            "experience": [
                {
                    "title": e.get("title"),
                    "company": e.get("company"),
                    "start_date": e.get("start_date"),
                    "end_date": e.get("end_date"),
                    "bullets": (e.get("bullets") or [])[:10],
                }
                for e in (profile.get("experience") or [])[:10]
            ],
            "education": [
                {
                    "degree": e.get("degree"),
                    "institution": e.get("institution"),
                    "graduation_year": e.get("graduation_year"),
                }
                for e in (profile.get("education") or [])[:4]
            ],
        }
        years_experience = _estimate_experience_years(clean_profile.get("experience") or [])
        requires_senior_portfolio = years_experience > 8
        prioritize_projects = _top_relevant_projects(clean_profile, jd, limit=4)
        project_focus_hint = (
            "Include a dedicated '## Key Projects' section with 3-4 deep-dive entries near the top (after Skills)."
            if years_experience > 8
            else "Keep projects concise: include only 1-2 tightly relevant project highlights and emphasize skills evidence."
            if years_experience < 3
            else "Include 2-3 project highlights tied directly to the target JD priorities."
        )

        prompt = (
            "You are an expert US resume writer specializing in ATS-optimized resumes for tech professionals.\n\n"
            "TASK: Generate a tailored ATS resume + story-based cover letter as STRICT JSON.\n\n"
            "STRICT RULES FOR RESUME (COGNITIVE INTELLIGENCE):\n"
            "1. ATS-safe Markdown only: NO tables, NO columns, NO images.\n"
            "2. Section ORDER: # CandidateName → contact line → ## SUMMARY → ## SKILLS → ## EXPERIENCE → ## EDUCATION\n"
            "3. Every experience entry uses ### Role | Company | Date range format + bullet points.\n"
            "4. Skills: include ALL candidate skills + any JD skills the candidate plausibly has.\n"
            "5. Do NOT invent employers, degrees, or specific metrics. Use 'improved' without inventing numbers.\n\n"
            "9. Strict No-URL Policy: Remove all URLs from body sections. URLs are allowed only in the header contact line.\n"
            "10. Experience-Aware Scaling: If total experience >8 years, MUST include '## Key Projects' with 3-4 deep-dive entries; if <3 years, include only 1-2 projects and focus on skills.\n"
            "11. Professional Narrative Rule: STRICTLY PROHIBITED from STAR-style headers/labels like 'Situation:', 'Task:', 'Action:', 'Result:'. Use strong action verbs (Spearheaded, Optimized, Architected, Led, Built) in natural narrative bullets.\n"
            "12. Contextual Synthesis Rule: Cross-reference Master Resume with JD and pull historically relevant projects to the top when JD priorities match (e.g., fraud detection experience from earlier roles).\n\n"
            "STRICT RULES FOR COVER LETTER:\n"
            "6. Story structure: Opening → Action paragraph → Result paragraph → Challenge/growth paragraph → Close\n"
            "7. Pick ONE strong project from experience and map to JD's biggest technical challenge.\n"
            "8. First person, professional US tone, 3-4 paragraphs.\n\n"
            "OUTPUT: Return ONLY a JSON object — no markdown fences, no other text:\n"
            "{\"resume_md\": \"...\", \"cover_md\": \"...\"}\n\n"
            f"CANDIDATE_PROFILE:\n{json.dumps(clean_profile, indent=2)}\n\n"
            f"EXPERIENCE_YEARS_ESTIMATE: {years_experience}\n"
            f"REQUIRES_SENIOR_PORTFOLIO: {requires_senior_portfolio}\n\n"
            f"PROJECT_FOCUS_DIRECTIVE: {project_focus_hint}\n\n"
            f"PRIORITIZED_RELEVANT_PROJECTS: {json.dumps(prioritize_projects, indent=2)}\n\n"
            f"TARGET_ROLE: {title}\n\n"
            f"JOB_DESCRIPTION:\n{jd[:8000]}\n"
        )

        j = self.gemini.generate_json(prompt, temperature=0.15, max_tokens=4000)
        if not isinstance(j, dict):
            return fallback_resume, fallback_cover

        resume_md = str(j.get("resume_md") or "").strip()
        cover_md = str(j.get("cover_md") or "").strip()

        # Validate output has required sections
        if not resume_md or not re.search(r"^##\s+", resume_md, re.M):
            resume_md = fallback_resume
        if not cover_md or len(cover_md) < 80:
            cover_md = fallback_cover

        resume_md = _sanitize_resume_markdown(resume_md, profile=clean_profile, title=title, jd=jd)
        resume_md = _strip_urls_except_header_contact(resume_md)
        cover_md = _sanitize_cover_letter(cover_md, profile=clean_profile, title=title)
        cover_md = _strip_urls_except_header_contact(cover_md)
        return resume_md, cover_md


def _estimate_experience_years(experience: List[Dict[str, Any]]) -> int:
    years = 0
    for e in experience or []:
        start = str(e.get("start_date") or "")
        end = str(e.get("end_date") or "")
        start_year = int(start[:4]) if re.match(r"^\d{4}", start) else None
        end_year = int(end[:4]) if re.match(r"^\d{4}", end) else None
        if start_year:
            years += max(1, (end_year or datetime.now().year) - start_year)
    return min(40, years)


def _strip_urls_except_header_contact(text: str) -> str:
    lines = (text or "").splitlines()
    out: List[str] = []
    header_done = False
    for ln in lines:
        stripped = ln.strip()
        if stripped.startswith("## "):
            header_done = True
        if header_done:
            ln = re.sub(r"https?://\S+", "", ln)
            ln = re.sub(r"\bwww\.\S+", "", ln)
        out.append(re.sub(r"\s{2,}", " ", ln).rstrip())
    return "\n".join(out).strip() + "\n"



def _sanitize_resume_markdown(resume_md: str, *, profile: Dict[str, Any], title: str, jd: str) -> str:
    txt = (resume_md or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.rstrip() for ln in txt.splitlines()]

    # Remove duplicate contact/summary labels often caused by malformed prompts.
    deduped: List[str] = []
    seen_compact: set[str] = set()
    for line in lines:
        compact = re.sub(r"\s+", " ", line.strip().lower())
        if compact and compact in seen_compact and (
            "professional summary" in compact or "http" in compact or "@" in compact
        ):
            continue
        if compact:
            seen_compact.add(compact)
        deduped.append(line)

    txt = "\n".join(deduped).strip()

    # Ban STAR headers and similar AI-template labels.
    txt = re.sub(r"^\s*(Situation|Task|Action|Result)\s*:\s*", "", txt, flags=re.I | re.M)

    # Ensure clean section headers with ATS-friendly title casing.
    txt = re.sub(r"^##\s*professional summary\s*$", "## Summary", txt, flags=re.I | re.M)
    txt = re.sub(r"^##\s*technical skills\s*$", "## Skills", txt, flags=re.I | re.M)

    required = ["## Summary", "## Skills", "## Experience", "## Education"]
    if not txt.startswith("# ") or any(sec.lower() not in txt.lower() for sec in required):
        return _build_fallback_resume(profile, title, jd)

    # Remove markdown emphasis from bullet labels to keep ATS parsing stable.
    txt = re.sub(r"\*\*(AI/ML|Cloud Engineering|Data Ops):\*\*", r"\1:", txt)
    return txt + ("\n" if not txt.endswith("\n") else "")


def _sanitize_cover_letter(cover_md: str, *, profile: Dict[str, Any], title: str) -> str:
    txt = (cover_md or "").strip()
    if txt.lower().startswith("cover letter"):
        txt = re.sub(r"^cover letter\s*", "", txt, flags=re.I).strip()

    # Ensure proper salutation and close.
    if "dear " not in txt.lower()[:40]:
        txt = f"Dear Hiring Manager,\n\n{txt}"
    if "sincerely" not in txt.lower():
        name = str(profile.get("name") or "Candidate")
        txt = txt.rstrip() + f"\n\nSincerely,\n{name}"

    # Guardrail against ultra-short cover letters.
    if len(txt) < 220:
        return _build_fallback_cover(profile, title)
    return txt + ("\n" if not txt.endswith("\n") else "")


def _strip_markdown(text: str) -> str:
    txt = text or ""
    txt = re.sub(r"```.*?```", " ", txt, flags=re.S)
    txt = re.sub(r"`([^`]*)`", r"\1", txt)
    txt = re.sub(r"!\[[^\]]*\]\([^)]*\)", " ", txt)
    txt = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", txt)
    txt = re.sub(r"^\s{0,3}#{1,6}\s*", "", txt, flags=re.M)
    txt = re.sub(r"^\s*[-*+]\s+", "", txt, flags=re.M)
    txt = re.sub(r"^\s*\d+\.\s+", "", txt, flags=re.M)
    txt = re.sub(r"[*_~]", "", txt)
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt


def _build_fallback_resume(profile: Dict[str, Any], title: str, jd: str) -> str:
    name = str(profile.get("name") or "Candidate Name")
    contact = profile.get("contact") or {}
    c_line = " | ".join([v for v in [contact.get("email"), contact.get("phone"), contact.get("location")] if v])
    if not c_line:
        c_line = "email@example.com | +1 (555) 555-5555"
    summary = str(profile.get("summary") or f"Results-driven professional targeting {title} roles.")
    skills = profile.get("skills") or extract_skills(jd)
    skills_block = "\n".join(f"- {s}" for s in skills[:20])

    exp_lines = []
    relevant_projects = _top_relevant_projects(profile, jd, limit=4)
    years_experience = _estimate_experience_years(profile.get("experience") or [])
    project_target = 4 if years_experience > 8 else 2 if years_experience < 3 else 3

    for e in (profile.get("experience") or [])[:5]:
        head = " | ".join([x for x in [e.get("title"), e.get("company"), f"{e.get('start_date') or ''} - {e.get('end_date') or 'Present'}"] if x])
        exp_lines.append(f"### {head}" if head else "### Experience")
        bullets = e.get("bullets") or ["Delivered measurable impact across core initiatives."]
        exp_lines.extend([f"- {b}" for b in bullets[:4]])
    if not exp_lines:
        exp_lines = ["### Relevant Experience", "- Delivered high-impact projects aligned to business goals."]

    edu_lines = []
    for e in (profile.get("education") or [])[:4]:
        row = " | ".join([x for x in [e.get("degree"), e.get("institution"), str(e.get("graduation_year") or "")] if x])
        if row:
            edu_lines.append(f"- {row}")
    if not edu_lines:
        edu_lines = ["- Education details available upon request."]

    project_lines: List[str] = []
    for idx, p in enumerate(relevant_projects[:project_target], start=1):
        project_header = " | ".join([x for x in [p.get("title") or f"Project {idx}", p.get("company") or ""] if x])
        project_lines.append(f"### {project_header}".rstrip())
        for bullet in (p.get("bullets") or ["Optimized delivery outcomes aligned to target role requirements."])[:3]:
            project_lines.append(f"- {bullet}")
    if not project_lines:
        project_lines = ["### Project Spotlight", "- Architected solutions aligned to the role requirements and business goals."]

    return (
        f"# {name}\n"
        f"{c_line}\n\n"
        f"## Summary\n{summary}\n\n"
        f"## Skills\n{skills_block}\n\n"
        "## Key Projects\n"
        + "\n".join(project_lines)
        + "\n\n"
        "## Experience\n"
        + "\n".join(exp_lines)
        + "\n\n## Education\n"
        + "\n".join(edu_lines)
        + "\n"
    )


def _build_fallback_cover(profile: Dict[str, Any], title: str) -> str:
    name = str(profile.get("name") or "Candidate")
    summary = str(profile.get("summary") or "I bring a strong track record of execution and collaboration.")
    return (
        f"Dear Hiring Manager,\n\n"
        f"I am excited to apply for the {title} opportunity. {summary}\n\n"
        "In my recent work, I partnered with cross-functional teams to deliver practical, scalable outcomes while balancing speed and quality. "
        "I enjoy turning ambiguous requirements into reliable systems that create user and business value.\n\n"
        "I would welcome the chance to discuss how my background aligns with your needs. Thank you for your time and consideration.\n\n"
        f"Sincerely,\n{name}"
    )
